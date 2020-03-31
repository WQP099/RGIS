import sys
import datetime
import os
import operator
import folium
import cv2
import time
import base64
import threading
import re
import webbrowser as wb
if hasattr(sys, 'frozen'):
    os.environ['PATH'] = sys._MEIPASS + ";" + os.environ['PATH']
from PyQt5 import QtWidgets

from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from collections import Counter

from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtGui import   QPainter ,QPixmap,QPen,QColor
from PyQt5.QtWidgets import (QGraphicsItem, QGraphicsObject, QGraphicsScene, QGraphicsView, QMainWindow)
from PyQt5.QtWidgets import QListWidgetItem, QCheckBox
from PyQt5.QtWebEngineWidgets import QWebEngineView

from PyQt5.QtWidgets import QApplication,QMainWindow,QFileDialog,QMenu,QLineEdit,QMessageBox,QColorDialog,QDesktopWidget
from PyQt5.QtCore import pyqtSignal,Qt,QRectF,QPoint,QPointF
import json
import tkinter as tk

from tkinter import filedialog
from use_sqlite import sqlitefun
from basicsfunction import basicsfun
from PyQt5.QtCore import QUrl
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *
from PyQt5.QtWebEngineWidgets import *

#导入区块链文件
from requests_post import RequestsPost

from UI.Ui_RGISMainActivity import Ui_RGISMain  # 使用vscode生成的调用方法
from UI.Ui_create_project import Ui_creatProject 
from UI.Ui_open_project import Ui_open_project
from UI.set_panel_info import Ui_setPanelInfo
from UI.Ui_num_dialog import Ui_numDialog
from UI.Ui_row_column_dialog import  Ui_row_Column_Dialog
from UI.Ui_update_panel_info import Ui_Upadte_Panel_info
from UI.Update_Number_Row import Ui_Update_number_Dialog
from UI.update_select_panel import Ui_update_Panel
from UI.update_panel_mode import Ui_Update_Panel_Model
from UI.upload_chain import Ui_Chain_object
from UI.Ui_Loading_object import Ui_Loading_object


class MyRGISMain(QMainWindow,Ui_RGISMain):
    signalMousePos = pyqtSignal(list,list,int)

    def __init__(self,parent = None):
        super(MyRGISMain, self).__init__(parent)
        self.setupUi(self)
        self.setWindowTitle("RGIS-地理信息标注系统v0001-20191219")

        self.rownum = 2 #单行标注点击次数
        self.row_column = 4 #多行点击次数
        self.rowvalue1=[] #单行第一次点击的点
        self.rowvalue2=[] #单行第二次点击的点
        self.rowcolumn = [[0,0],[0,0],[0,0],[0,0]] #多行第一次点击的点
        # self.rowcolumn2 = [] #多行第二次点击的点
        # self.rowcolumn3 = [] #多行第三次点击的点
        # self.rowcolumn4 = [] #多行第四次点击的点
        self.movemode = 0 # 0表示固定 1 表示移动
        self.MapImage.setDragMode(QtWidgets.QGraphicsView.NoDrag)
        self.cb_label_mode.addItems(["查看","单点标注","单行标注","多行标注"])
        self.cb_label_mode.setEnabled(False)
        self.cb_label_mode.currentIndexChanged.connect(self.model_currentChanged)
        self.bool_getpanel_wh = 0  # 获取组件长宽
        self.projecttime = ""
        self.panel_info_all_list = []  # 原始信息列表 保持和数据库中的一直
        self.panel_info_dictlist = []  # 面板全部信息  用列表字典保存
        self.general_panel_showlist = []  # 一般显示列表  通常全部组件是一种颜色 用绿色表示
        self.clicked_panel_showlist = [] # 选中组件显示列表  一般表示多个或者一个选中的显示 颜色用红色表示
        self.record_GPS_list = [] #记录上次的GPS信息

        self.cb_panel_mode.setEnabled(False)


        self.bool_region_move_panel = 0 #用于控制是否允许移动组件
        self.bool_batch_update_panel = 0 #用于控制是否批量修改组件
        self.bool_batch_del_panel = 0 #用于控制是否批量删除组件
        self.bool_updage_number = 0

        self.selectItem = []
        #初始化三个显示颜色面板
        self.redboxlist = []
        self.greenboxlist = []
        self.blueboxlist = []
        self.pos_check = []


        self.action_create_project.triggered.connect(self.create_project) # 创建GIS项目
        self.action_history_project.triggered.connect(self.history_project) # 打开历史项目
        self.action_export_GPS_image.triggered.connect(self.export_gps_image)#导出GPS
        self.action_import_GPS_image.triggered.connect(self.import_gps_image)  # 导入GPS
        self.action_gd_map.triggered.connect(self.open_gd_map)#打开地图
        self.action_export_word.triggered.connect(self.export_word)#导出word
        self.action_chain_table.triggered.connect(self.chain_table)#区块链表
        self.menu77.setEnabled(False)

        self.but_zoomout.clicked.connect(self.MapimageZoonOut) # 放大按钮
        self.but_zoomin.clicked.connect(self.MapimageZoonIn) # 缩小按钮
        self.but_rest.clicked.connect(self.MapimageZoonRest)  # 刷新按钮
        self.but_movemode.clicked.connect(self.changeMoveMode) # 改变点击模式 切换拖动和点击
        self.but_getpanel_wh.clicked.connect(self.getpanel_wh) # 获取组件宽高
        self.but_setpanel_info.clicked.connect(self.setpanelInfo) # 设置组件信息
        self.btn_del_panel.clicked.connect(self.del_panel_info)#删除信息
        self.btn_select_panel_info.clicked.connect(self.update_panel_info)        #修改信息
        self.but_update_panel_mode.clicked.connect(self.update_panel_mode) #修改模板信息
        self.but_delete_panel_mode.clicked.connect(self.delete_panel_mode)#删除模板信息

        self.listPanelID.setContextMenuPolicy(Qt.CustomContextMenu)
        self.listPanelID.customContextMenuRequested.connect(self.select_right_menu)#绑定列表
        #self.listPanelID.itemDoubleClicked.connect(self.updateItem)  # 设置条目点击事件
        self.listPanelID.itemClicked.connect(self.select_panel_item)  # 条目单点事件

        self.btn_move_no.setVisible(False)
        self.btn_move_ok.setVisible(False)
        self.btn_move_ok.clicked.connect(self.start_fun_ok)#确定开始功能事件
        self.btn_move_no.clicked.connect(self.end_fun_no)#取消功能事件
        self.btn_move_panel.clicked.connect(self.region_loc_alter)#多面板位置改动

        self.btn_update_number.clicked.connect(self.update_number)#重新编号
        self.btn_update_number_no.setVisible(False)
        self.btn_update_number_ok.setVisible(False)
        self.btn_update_number_no.clicked.connect(self.update_number_No)#取消重新编号
        self.btn_update_number_ok.clicked.connect(self.update_number_Ok)#确定重新编号
        self.btn_GPS_computing.clicked.connect(self.GPS_computing)#GPS计算
        self.btn_GPS_range_computing.clicked.connect(self.GPS_range_computing)#GPS区域标注
        self.btn_GPS_range_computing_ok.setVisible(False)
        self.btn_GPS_range_computing_ok.clicked.connect(self.GPS_range_computing_OK)#GPS区域标注确定
        self.btn_GPS_range_computing_cancel.setVisible(False)
        self.btn_GPS_range_computing_cancel.clicked.connect(self.GPS_range_computing_Canel)  # GPS区域标注确定
        self.btn_GPS_undone_computing.clicked.connect(self.GPS_undone_computing)#GPS撤销
        self.btn_GPS_clear.clicked.connect(self.GPS_clear)#GPS清除

        self.btn_batch_update_no.setVisible(False)
        self.btn_batch_update_ok.setVisible(False)
        self.btn_batch_update_panel.clicked.connect(self.Batch_update_panel)#批量修改
        self.btn_batch_update_ok.clicked.connect(self.Batch_update_ok)#确认批量修改
        self.btn_batch_update_no.clicked.connect(self.Batch_update_no)#取消批量修改

        self.btn_batch_del_no.setVisible(False)
        self.btn_batch_del_ok.setVisible(False)
        self.btn_batch_del_panel.clicked.connect(self.Batch_del_panel)  # 批量删除
        self.btn_batch_del_ok.clicked.connect(self.Batch_del_ok)  # 确认批量删除
        self.btn_batch_del_no.clicked.connect(self.Batch_del_no)  # 取消批量删除
        root = tk.Tk()
        root.withdraw()
        self.select_pos=[]
        self.center()

        # win_Num_Dialog = Num_Dialog(self)
        # win_Num_Dialog.show()
    def center(self):
        size = self.geometry()
        screen=QDesktopWidget().screenGeometry()
        self.move((screen.width() - size.width()) /2,(screen.height() - size.height()) /2)

    def closeEvent(self, event):
        if self.projecttime!="":
            self.get_old_template(self.projecttime)
    #选择组件模板
    def templateClick(self):
        templateName = self.cb_panel_mode.currentText()
        if templateName == '':
            old_model = sqlitefun.sqlite_select_config("panel_model")
            templateName = old_model[0][0]
        print("模型名称", templateName)
        panel_json = sqlitefun.sqlite_select_template(templateName)
        panel_dict = json.loads(panel_json[0][0])
        sqlitefun.sqlite_insert_config("label_type", panel_dict["label_type"])
        sqlitefun.sqlite_insert_config("panel_model", panel_dict["panel_model"])
        sqlitefun.sqlite_insert_config("panel_power", panel_dict["panel_power"])
        sqlitefun.sqlite_insert_config("panel_wh", panel_dict["panel_wh"])
        sqlitefun.sqlite_insert_config("panel_altitude", panel_dict["panel_altitude"])
        sqlitefun.sqlite_insert_config("panel_angle", panel_dict["panel_angle"])
        sqlitefun.sqlite_insert_config("panel_manufactor", panel_dict["panel_manufactor"])
        sqlitefun.sqlite_insert_config("panel_arrange", panel_dict["panel_arrange"])
        sqlitefun.sqlite_insert_config("panel_color", panel_dict["panel_color"])
        mode_info = templateName+"模型内容信息为：\n"
        mode_info += "标注类型："+panel_dict["label_type"]+","
        mode_info += "组件型号：" + panel_dict["panel_model"]+","
        mode_info += "组件功率：" + panel_dict["panel_power"] + ","
        mode_info += "组件宽高：" + panel_dict["panel_wh"] + ","
        mode_info += "组件海拔：" + panel_dict["panel_altitude"] + ","
        mode_info += "组件角度：" + panel_dict["panel_angle"] + ","
        mode_info += "组件厂家：" + panel_dict["panel_manufactor"] + ","
        mode_info += "组件排布：" + panel_dict["panel_arrange"]
        self.te_bottom_hint.setText(mode_info)
    #选择模型后的点击事件
    def model_currentChanged(self):
        model = self.cb_label_mode.currentText()
        if model=="查看":
            self.but_setpanel_info.setEnabled(True)
            self.but_update_panel_mode.setEnabled(True)
            self.but_delete_panel_mode.setEnabled(True)
            self.but_getpanel_wh.setEnabled(True)
            self.btn_move_panel.setEnabled(True)
            self.btn_update_number.setEnabled(True)
            self.btn_GPS_computing.setEnabled(True)
            self.btn_GPS_range_computing.setEnabled(True)
            self.btn_GPS_undone_computing.setEnabled(True)
            self.btn_GPS_clear.setEnabled(True)
            self.btn_del_panel.setEnabled(True)
            self.btn_select_panel_info.setEnabled(True)
            self.btn_batch_update_panel.setEnabled(True)
            self.btn_batch_del_panel.setEnabled(True)
            self.cb_panel_mode.setEnabled(True)
        else:
            if len(self.select_pos)>0:
                self.delete_pos(self.select_pos)
                self.select_pos = []
            if model =="查看":
                self.row_column = 4
                self.rownum = 2
            if model =="单行标注":
                self.row_column = 4
            if model =="多行标注":
                self.rownum = 2
            self.te_bottom_hint.setText("当前为"+model+"状态")

            self.but_setpanel_info.setEnabled(False)
            self.but_update_panel_mode.setEnabled(False)
            self.but_delete_panel_mode.setEnabled(False)
            self.but_getpanel_wh.setEnabled(False)
            self.btn_move_panel.setEnabled(False)
            self.btn_update_number.setEnabled(False)
            self.btn_GPS_computing.setEnabled(False)
            self.btn_GPS_range_computing.setEnabled(False)
            self.btn_GPS_undone_computing.setEnabled(False)
            self.btn_GPS_clear.setEnabled(False)
            self.btn_del_panel.setEnabled(False)
            self.btn_select_panel_info.setEnabled(False)
            self.btn_batch_update_panel.setEnabled(False)
            self.btn_batch_del_panel.setEnabled(False)
            self.cb_panel_mode.setEnabled(False)
    #修改模板信息
    def update_panel_mode(self):
        print("修改模板信息")
        self.Update_Model = Update_Panel_Model_class(self)
        self.Update_Model.setT(1)
        self.Update_Model.Update_Panel_Model_Signal.connect(self.getDialogSignal)
        self.Update_Model.show()
    #删除模板信息
    def delete_panel_mode(self):
        print("删除模板信息")
        self.Update_Model = Update_Panel_Model_class(self)
        self.Update_Model.setT(2)
        self.Update_Model.Update_Panel_Model_Signal.connect(self.getDialogSignal)
        self.Update_Model.show()
    #区块链
    def chain_table(self):
        print("区块链信息")
        #获取当前项目的所有点信息，模板信息，图片名称，项目名称
        #项目名称
        projectlist = sqlitefun.sqlite_select_localist()
        token = RequestsPost.getToken()
        plant_info_request = RequestsPost.plant_chain_tabel_all_select(token)
        print("plant_info_request",plant_info_request)
        plantlist = []
        for plant_info in plant_info_request:
            print(plant_info)
            #plant_dict = json.loads(str(plant_info))
            plantlist.append(plant_info["plant_id"]+"/"+plant_info["chain_hand"])

        #获取项目名称
        upload_chain = Upload_Chain_class(self)
        upload_chain.setT(projectlist,plantlist)
        upload_chain.Upload_Chain_Signal.connect(self.getDialogSignal)
        upload_chain.show()
    #导出word
    def export_word(self):

        Folderpath = filedialog.asksaveasfilename(title=u'保存报告', filetypes=[("DOCX", ".docx")],
                                                  initialfile=self.projectName + "导出报告.docx")
        if Folderpath:
            print("导出word文档")
            document = Document()
            document.add_heading('中科利丰', 0)
            document.add_heading('1.全景图', level=1)
            self.doc_image()
            document.add_picture('doc_image.jpg', width=Inches(5))
            document.add_heading('2.故障统计结果',level=1)
            table = document.add_table(rows=1, cols=7)
            table.style="Light List Accent 1"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '组件厂家'
            hdr_cells[1].text = '组件型号'
            hdr_cells[2].text = '组串数量'
            hdr_cells[3].text = '组件数量'
            hdr_cells[4].text = '组件功率'
            hdr_cells[5].text = '总功率'
            hdr_cells[6].text = '组件颜色'
            #获取厂家，并去除重复
            Model = []
            for panel_info in self.panel_info_all_list:
                panel_dict = json.loads(panel_info[1])
                Model.append(panel_dict["panel_model"])
            new_Model = dict(Counter(Model))
            print("组件型号信息：", new_Model)
            #model_info_list=[]#厂家信息List
            for model_name,model_num in new_Model.items():
                panel_model = model_name #组件厂家
                panel_num = model_num #组件数量
                panel_manufactor=""#组件厂家
                panel_arrange_num=0#组串数量
                panel_power = ""#组件功率

                for panel_info in self.panel_info_all_list:
                    panel_dict = json.loads(panel_info[1])
                    if panel_model == panel_dict["panel_model"]:
                        panel_manufactor = panel_dict["panel_manufactor"]
                        panel_arrange = panel_dict["panel_arrange"]
                        panel_arrange_c = int(panel_arrange.split('x')[0])
                        panel_arrange_r = int(panel_arrange.split('x')[1])
                        panel_arrange_num+=int(panel_arrange_c*panel_arrange_r)
                        panel_power = panel_dict["panel_power"]
                        panel_color = panel_dict["panel_color"]
                print("panel_power",panel_power)
                colorstr=""
                panel_all_power = int(panel_power)*int(panel_arrange_num)#组件总功率
                row_cells = table.add_row().cells
                row_cells[0].text = panel_manufactor
                row_cells[1].text = panel_model
                row_cells[2].text = str(panel_num)
                row_cells[3].text = str(panel_arrange_num)
                row_cells[4].text = panel_power
                row_cells[5].text = str(panel_all_power)

                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value=panel_color))
                row_cells[6]._tc.get_or_add_tcPr().append(shading_elm_1)
                #row_cells[6].text = "123"
            document.save(Folderpath)

    def doc_image(self):
        panoramic_bin = sqlitefun.sqlite_select_imagebin(self.projecttime, "panoramic")[0][0]
        imagepath = "doc_image.jpg"
        with open(imagepath, 'wb') as f:
            f.write(panoramic_bin)
        img = cv2.imread("doc_image.jpg")
        for panel_info in self.panel_info_all_list:
            panel_dict = json.loads(panel_info[1])
            pos = panel_dict["panel_pos"]
            panel_wh = panel_dict["panel_wh"]
            panel_model = panel_dict["panel_model"]
            panel_w = int(panel_wh.split(',')[0])
            panel_h = int(panel_wh.split(',')[1])
            pos_x = int(pos.split(',')[0])
            pos_y = int(pos.split(',')[1])
            Lift_X = int(pos_x - (panel_w / 2))
            Lift_Y = int(pos_y - (panel_h / 2))
            Right_X = int(pos_x + (panel_w / 2))
            Right_Y = int(pos_y + (panel_h / 2))
            panel_color = panel_dict['panel_color']
            r = int(panel_color[1:3], 16)
            g = int(panel_color[3:5], 16)
            b = int(panel_color[5:7], 16)
            img = cv2.rectangle(img, (Lift_X, Lift_Y), (Right_X, Right_Y), (b, g, r), -1)  # 绿色

        cv2.imwrite('doc_image.jpg', img)
    #将颜色转为16进制m
    def RGB_to_Hex(self,tmp):
        rgb = tmp.split(',')  # 将RGB格式划分开来
        colorStr = ''
        for i in rgb:
            num = int(i)  # 将str转int
            # 将R、G、B分别转化为16进制拼接转换并大写
            colorStr += str(hex(num))[-2:].replace('x', '0').upper()
        return colorStr
    #批量删除
    def Batch_del_panel(self):
        if self.bool_updage_number==1:
            self.te_bottom_hint.setText("正在修改组件编号，不可进行批量修改")
        elif self.bool_region_move_panel == 1:
            self.te_bottom_hint.setText("正在移动组件，不可进行批量修改")
        else:
            if self.cb_label_mode.currentText() == "查看":
                self.btn_batch_del_ok.setVisible(True)
                self.btn_batch_del_no.setVisible(True)
                self.btn_batch_update_panel.setEnabled(False)
                self.btn_GPS_range_computing.setEnabled(False)
                self.btn_update_number.setEnabled(False)
                self.btn_move_panel.setEnabled(False)
                self.btn_batch_del_panel.setEnabled(False)
                self.btn_del_panel.setEnabled(False)
                self.btn_select_panel_info.setEnabled(False)
                self.MapImage.setDragMode(QGraphicsView.RubberBandDrag)  # 设置为可以拉框
                self.bool_batch_del_panel = 1
    def Batch_del_ok(self):
        #确认删除
        self.btn_move_panel.setEnabled(True)
        self.btn_batch_update_panel.setEnabled(True)
        self.btn_batch_del_panel.setEnabled(True)
        self.btn_GPS_range_computing.setEnabled(True)
        self.btn_update_number.setEnabled(True)
        self.btn_del_panel.setEnabled(True)
        self.btn_select_panel_info.setEnabled(True)
        self.btn_batch_del_ok.setVisible(False)  # 隐藏掉功能按钮的取消键
        self.btn_batch_del_no.setVisible(False)  # 隐藏掉功能按钮的确认键
        # 模式还原到使用之前
        if self.movemode == 1:
            self.MapImage.setDragMode(QtWidgets.QGraphicsView.ScrollHandDrag)
        else:
            self.MapImage.setDragMode(QtWidgets.QGraphicsView.NoDrag)
        # 获取选中的组件
        self.panel_info_all_list = sqlitefun.sqlite_select_panelinfolist(self.projecttime)
        selectBox = []
        for box in self.general_panel_showlist:
            if box.isSelected():
                for panel_info in self.panel_info_all_list:
                    boxkey = box.key
                    if boxkey == panel_info[0]:
                        sqlitefun.sqlit_delete_panelinfo_json(self.projecttime, boxkey)
                selectBox.append(boxkey)
                print("boxkey",boxkey)
        self.getlSignal_refresh([4, selectBox])
        self.bool_batch_del_panel = 0
        # if len(selectBox) > 0:
        #     for selectPanel in selectBox:
        #         for panel_info in self.panel_info_all_list:
        #             panel_info_dict = json.loads(panel_info[1])  # 转为字典
        #             if selectPanel == panel_info[0]:
        #                 panelId = panel_info[0]
        #                 print("删除组件 ",panelId)
        #                 sqlitefun.sqlit_delete_panelinfo_json(self.projecttime, panelId)
        #                 #self.panel_info_all_list.remove(panel_info)


    def Batch_del_no(self):
        #取消删除
        self.bool_batch_del_panel = 0
        self.btn_move_panel.setEnabled(True)
        self.btn_batch_update_panel.setEnabled(True)
        self.btn_batch_del_panel.setEnabled(True)
        self.btn_GPS_range_computing.setEnabled(True)
        self.btn_update_number.setEnabled(True)
        self.btn_del_panel.setEnabled(True)
        self.btn_select_panel_info.setEnabled(True)
        self.btn_batch_del_ok.setVisible(False)  # 隐藏掉功能按钮的取消键
        self.btn_batch_del_no.setVisible(False)  # 隐藏掉功能按钮的确认键
        # 模式还原到使用之前
        if self.movemode == 1:
            self.MapImage.setDragMode(QtWidgets.QGraphicsView.ScrollHandDrag)
        else:
            self.MapImage.setDragMode(QtWidgets.QGraphicsView.NoDrag)
    #批量修改
    def Batch_update_panel(self):
        if self.bool_updage_number==1:
            self.te_bottom_hint.setText("正在修改组件编号，不可进行批量修改")
        elif self.bool_region_move_panel == 1:
            self.te_bottom_hint.setText("正在移动组件，不可进行批量修改")
        else:
            if self.cb_label_mode.currentText() == "查看":
                print("多面板位置修改")
                self.btn_batch_update_ok.setVisible(True)
                self.btn_batch_update_no.setVisible(True)
                self.btn_batch_update_panel.setEnabled(False)
                self.btn_GPS_range_computing.setEnabled(False)
                self.btn_update_number.setEnabled(False)
                self.btn_move_panel.setEnabled(False)
                self.btn_batch_del_panel.setEnabled(False)
                self.btn_select_panel_info.setEnabled(False)
                self.btn_del_panel.setEnabled(False)
                self.MapImage.setDragMode(QGraphicsView.RubberBandDrag)  # 设置为可以拉框
                self.bool_batch_update_panel = 1
                #if self.bool_batch_update_panel ==0:
                    #self.bool_batch_update_panel=1 #设置为1的时候  允许移动组件

                    #self.btn_move_panel.setEnabled(False) #设置为不可点击
                    # for box in self.general_panel_showlist:
                    #     box.bool_IsMovable = 1
                    #     box.setFlag(QGraphicsItem.ItemIsMovable,True)
            else:
                QMessageBox.information(self, "消息", "请将操作状态改为查看才可批量修改组件", QMessageBox.Yes)

    def Batch_update_ok(self):
        #确认批量修改
        self.btn_move_panel.setEnabled(True)
        self.btn_batch_update_panel.setEnabled(True)
        self.btn_batch_del_panel.setEnabled(True)
        self.btn_GPS_range_computing.setEnabled(True)
        self.btn_update_number.setEnabled(True)
        self.btn_select_panel_info.setEnabled(True)
        self.btn_del_panel.setEnabled(True)
        self.btn_batch_update_ok.setVisible(False)  # 隐藏掉功能按钮的取消键
        self.btn_batch_update_no.setVisible(False)  # 隐藏掉功能按钮的确认键
        # 模式还原到使用之前
        if self.movemode == 1:
            self.MapImage.setDragMode(QtWidgets.QGraphicsView.ScrollHandDrag)
        else:
            self.MapImage.setDragMode(QtWidgets.QGraphicsView.NoDrag)
        #获取选中的组件
        selectBox = []
        select_panel_info_dict=[]
        for box in self.general_panel_showlist:
            if box.isSelected():
                boxkey = box.key
                selectBox.append(boxkey)
        if len(selectBox)>0:
            for panel_info in self.panel_info_all_list:
                panel_info_dict = json.loads(panel_info[1])  # 转为字典
                if selectBox[0] in panel_info[0]:
                    select_panel_info_dict = panel_info_dict
            # 弹出信息修改框
            self.update_select_panel = Select_Update_Panel_info()
            self.update_select_panel.set_batch_update(selectBox)
            self.update_select_panel.setT(select_panel_info_dict, None,self.listPanelID)
            self.update_select_panel.Select_Update_Panel_info_Signal.connect(self.getDialogSignal)
            self.update_select_panel.show()



    def Batch_update_no(self):
        #取消批量修改
        # 关闭区域选中 修改位置
        self.btn_move_panel.setEnabled(True)
        self.btn_batch_update_panel.setEnabled(True)
        self.btn_batch_del_panel.setEnabled(True)
        self.btn_GPS_range_computing.setEnabled(True)
        self.btn_update_number.setEnabled(True)
        self.btn_select_panel_info.setEnabled(True)
        self.btn_del_panel.setEnabled(True)
        self.btn_batch_update_ok.setVisible(False)  # 隐藏掉功能按钮的取消键
        self.btn_batch_update_no.setVisible(False)  # 隐藏掉功能按钮的确认键
        # 模式还原到使用之前
        if self.movemode == 1:
            self.MapImage.setDragMode(QtWidgets.QGraphicsView.ScrollHandDrag)
        else:
            self.MapImage.setDragMode(QtWidgets.QGraphicsView.NoDrag)
        self.bool_batch_update_panel = 0
    def open_gd_map(self):
        # 转换经纬度
        gd_gps_list = []
        for panelInfo in self.panel_info_all_list:
            dict_info = json.loads(panelInfo[1])
            gps_w = dict_info["panel_gps"].split(',')[1]  # 纬度
            gps_j = dict_info["panel_gps"].split(',')[0]  # 经度
            lnglat = basicsfun.wgs84togcj02(float(gps_j), float(gps_w))
            lnglat_list = list(lnglat)
            gps = str(round(lnglat[0], 6)) + "," + str(round(lnglat[1], 6))
            dataItem = {'lnglat': list(lnglat_list), 'panel_id': panelInfo[0]}
            gd_gps_list.append(dataItem)
        #print("debug ", gd_gps_list)
        print("打开地图",gd_gps_list[0]['lnglat'])
        map = folium.Map(
            location=gd_gps_list[0]['lnglat'],
            #location=[41.22837816823197,105.81928501485979],
            zoom_start=8,
            #tiles='http://webrd02.is.autonavi.com/appmaptile?lang=zh_cn&size=1&scale=1&style=7&x={x}&y={y}&z={z}',
            tiles='https://www.google.cn/maps/vt?lyrs=s@189&gl=cn&x={x}&y={y}&z={z}',
            attr="google地图"
        )
        for gd_gps in gd_gps_list:
            gps = gd_gps['lnglat']
            title_info = gd_gps['panel_id']
            folium.Marker(gps, popup='<b>'+title_info+'</b>',icon=folium.Icon(color='red')).add_to(map)
        map.save('gd_map.html')
        boo = basicsfun.update_html('gd_map.html')
        #boo = True
        if boo:
            current_work_dir = os.getcwd()
            print(os.getcwd())  # 获得当前目录
            print("current",current_work_dir)
            self.gdwindow = GD_map_Window()
            file_path = current_work_dir +"\\gd_map.html"
            file_path = file_path.replace('\\', '/')
            print("文件地址",file_path)
            file_url = QUrl(file_path)
            self.gdwindow.setUrl(file_url)
            print("地图html ",file_url)
            self.gdwindow.show()
        else:
            QMessageBox.information(self, "消息", "地图打开失败", QMessageBox.Yes)

    #GPS计算
    def GPS_computing(self):
        if self.bool_updage_number==1:
            self.te_bottom_hint.setText("正在修改组件编号，不可进行GPS计算")
        elif self.bool_batch_update_panel == 1:
            self.te_bottom_hint.setText("正在批量修改组件信息，不可进行GPS计算")
        elif self.bool_region_move_panel == 1:
            self.te_bottom_hint.setText("正在移动组件，不可进行GPS计算")
        elif self.bool_batch_del_panel ==1:
            self.te_bottom_hint.setText("正在批量删除，不可进行GPS计算")
        else:
            #GPS计算
            print("GPS计算")
            inputxy = []
            inputgps = []
            inputpos = []
            #获取4个点的值
            select_xy_list = []
            # GPS_list_value = ["2,2","4,2","4,4","2,4"]
            GPS_list_value = []
            select_pos_value = []
            self.record_GPS_list=[]
            for select_itme in self.panel_info_all_list:
                panel_info_dict = json.loads(select_itme[1])  # 转为字典
                if "1,1" != panel_info_dict["panel_gps"]:
                    select_xy_list.append(panel_info_dict["panel_pos"])
                    GPS_list_value.append(panel_info_dict["panel_gps"])
                GPS_info = select_itme[0]+":"+panel_info_dict["panel_gps"]
                self.record_GPS_list.append(GPS_info)
                select_pos_value.append(panel_info_dict["panel_pos"])
            if len(GPS_list_value)<4:
                QMessageBox.information(self, "消息", "最少需要4个GPS点信息", QMessageBox.Yes)
            else:
                for inputxyliststr in select_xy_list:
                    inputxy.append(inputxyliststr.split(","))

                for inputgpsliststr in GPS_list_value:
                    inputgps.append(inputgpsliststr.split(","))

                for inputposliststr in select_pos_value:
                    inputpos.append(inputposliststr.split(","))
                res = basicsfun.get_transform_mat(inputxy,inputgps,inputpos)
                print("GPS    ",res)
                gps_list = res.split(' ')
                info_all_list = []
                for i in range(len(gps_list)):
                    panel_item = self.panel_info_all_list[i]
                    panel_dict = json.loads(panel_item[1])
                    panel_dict["panel_gps"] = gps_list[i]
                    panel_info_json = json.dumps(panel_dict, ensure_ascii=False)
                    info_all_list.append([panel_item[0],panel_info_json])
                sqlitefun.sqlite_insert_panelinfo_all_json(self.projecttime, info_all_list)
                self.getlSignal_refresh([2])
                self.te_bottom_hint.setText("GPS计算完成。")

    def GPS_range_computing(self):
        print("GPS区域计算")
        if self.bool_updage_number==1:
            self.te_bottom_hint.setText("正在修改组件编号，不可进行批量修改")
        elif self.bool_region_move_panel == 1:
            self.te_bottom_hint.setText("正在移动组件，不可进行批量修改")
        else:
            if self.cb_label_mode.currentText() == "查看":
                self.btn_GPS_range_computing_ok.setVisible(True)
                self.btn_GPS_range_computing_cancel.setVisible(True)
                self.btn_GPS_range_computing.setEnabled(False)
                self.btn_batch_update_panel.setEnabled(False)
                self.btn_move_panel.setEnabled(False)
                self.btn_batch_del_panel.setEnabled(False)
                self.btn_update_number.setEnabled(False)
                self.btn_GPS_computing.setEnabled(False)
                self.MapImage.setDragMode(QGraphicsView.RubberBandDrag)  # 设置为可以拉框
                #self.bool_batch_del_panel = 1
    def GPS_range_computing_OK(self):
        print("GPS区域确定")
        self.btn_GPS_range_computing_ok.setVisible(False)
        self.btn_GPS_range_computing_cancel.setVisible(False)
        self.btn_GPS_range_computing.setEnabled(True)
        self.btn_batch_update_panel.setEnabled(True)
        self.btn_move_panel.setEnabled(True)
        self.btn_batch_del_panel.setEnabled(True)
        self.btn_update_number.setEnabled(True)
        self.btn_GPS_computing.setEnabled(True)
        # 模式还原到使用之前
        if self.movemode == 1:
            self.MapImage.setDragMode(QtWidgets.QGraphicsView.ScrollHandDrag)
        else:
            self.MapImage.setDragMode(QtWidgets.QGraphicsView.NoDrag)
        # 获取选中的组件
        self.panel_info_all_list = sqlitefun.sqlite_select_panelinfolist(self.projecttime)
        selectBox = []
        inputxy = []
        inputgps = []
        inputpos = []
        # 获取4个点的值
        select_xy_list = []
        GPS_list_value = []
        select_pos_value = []
        self.record_GPS_list=[]
        for box in self.general_panel_showlist:
            if box.isSelected():
                for panel_info in self.panel_info_all_list:
                    boxkey = box.key
                    if boxkey == panel_info[0]:
                        selectBox.append(panel_info)
                print("boxkey", boxkey)
        for selectGPS in selectBox:
            panel_info_dict = json.loads(selectGPS[1])  # 转为字典
            if "1,1" != panel_info_dict["panel_gps"]:
                select_xy_list.append(panel_info_dict["panel_pos"])
                GPS_list_value.append(panel_info_dict["panel_gps"])
            GPS_info = selectGPS[0] + ":" + panel_info_dict["panel_gps"]
            self.record_GPS_list.append(GPS_info)
            select_pos_value.append(panel_info_dict["panel_pos"])
        if len(GPS_list_value) < 4:
            QMessageBox.information(self, "消息", "最少需要4个GPS点信息", QMessageBox.Yes)
        else:
            for inputxyliststr in select_xy_list:
                inputxy.append(inputxyliststr.split(","))

            for inputgpsliststr in GPS_list_value:
                inputgps.append(inputgpsliststr.split(","))

            for inputposliststr in select_pos_value:
                inputpos.append(inputposliststr.split(","))
            res = basicsfun.get_transform_mat(inputxy, inputgps, inputpos)
            print("区域修改GPS    ", res)
            gps_list = res.split(' ')
            info_all_list = []
            for i in range(len(gps_list)):
                panel_item = selectBox[i]
                panel_dict = json.loads(panel_item[1])
                panel_dict["panel_gps"] = gps_list[i]
                panel_info_json = json.dumps(panel_dict, ensure_ascii=False)
                info_all_list.append([panel_item[0], panel_info_json])
            sqlitefun.sqlite_insert_panelinfo_all_json(self.projecttime, info_all_list)
            self.getlSignal_refresh([2])
            self.te_bottom_hint.setText("GPS区域修改已完成。")

    def GPS_range_computing_Canel(self):
        print("GPS区域取消")
        self.btn_GPS_range_computing_ok.setVisible(False)
        self.btn_GPS_range_computing_cancel.setVisible(False)
        self.btn_GPS_range_computing.setEnabled(True)
        self.btn_batch_update_panel.setEnabled(True)
        self.btn_move_panel.setEnabled(True)
        self.btn_batch_del_panel.setEnabled(True)
        self.btn_update_number.setEnabled(True)
        self.btn_GPS_computing.setEnabled(True)
        # 模式还原到使用之前
        if self.movemode == 1:
            self.MapImage.setDragMode(QtWidgets.QGraphicsView.ScrollHandDrag)
        else:
            self.MapImage.setDragMode(QtWidgets.QGraphicsView.NoDrag)

    def GPS_undone_computing(self):
        print("GPS 撤消操作")
        info_all_list=[]
        #将上次记录的GPS值给赋值回去
        for record_GPS in self.record_GPS_list:
            panel_id = record_GPS.split(':')[0]
            panel_GPS = record_GPS.split(':')[1]
            for panel_info in self.panel_info_all_list:
                panel_dict = json.loads(panel_info[1])
                if panel_id==panel_info[0]:
                    panel_dict["panel_gps"] = panel_GPS
                    panel_info_json = json.dumps(panel_dict, ensure_ascii=False)
                    info_all_list.append([panel_info[0], panel_info_json])

        sqlitefun.sqlite_update_panelinfo_all_json(self.projecttime, info_all_list)
        self.getlSignal_refresh([2])
        self.te_bottom_hint.setText("GPS撤销完成。")

    #GPS清除
    def GPS_clear(self):
        if self.bool_updage_number==1:
            self.te_bottom_hint.setText("正在修改组件编号，不可进行GPS清除")
        elif self.bool_batch_update_panel == 1:
            self.te_bottom_hint.setText("正在批量修改组件信息，不可进行GPS清除")
        elif self.bool_region_move_panel == 1:
            self.te_bottom_hint.setText("正在移动组件，不可进行GPS清除")
        elif self.bool_batch_del_panel ==1:
            self.te_bottom_hint.setText("正在批量删除，不可进行GPS清除")
        else:
            reply = QMessageBox.question(self, "GPS清除提示", "你确定要清除掉所有GPS信息?",
                                         QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.Yes:
                print("GPS清除")
                info_all_list=[]
                for i in range(len(self.panel_info_all_list)):
                    panel_item = self.panel_info_all_list[i]
                    panel_dict = json.loads(panel_item[1])
                    panel_dict["panel_gps"] = "1,1"
                    panel_info_json = json.dumps(panel_dict,ensure_ascii=False)
                    info_all_list.append([panel_item[0], panel_info_json])
                sqlitefun.sqlite_insert_panelinfo_all_json(self.projecttime, info_all_list)
                self.getlSignal_refresh([2])
                self.te_bottom_hint.setText("GPS清除完成。")
    #重新编号
    def update_number(self):
        print("item  重新编号")
        if self.cb_label_mode.currentText() == "查看":

            if self.bool_region_move_panel ==1:
                self.te_bottom_hint.setText("正在移动组件，不可进行组件重新编号")
            elif self.bool_batch_update_panel == 1:
                self.te_bottom_hint.setText("正在批量修改组件信息，不可进行组件重新编号")
            elif self.bool_batch_del_panel == 1:
                self.te_bottom_hint.setText("正在批量删除，不可进行组件重新编号")
            else:
                self.btn_move_panel.setEnabled(False)
                self.btn_batch_update_panel.setEnabled(False)
                self.btn_batch_del_panel.setEnabled(False)
                self.btn_GPS_range_computing.setEnabled(False)
                self.btn_update_number.setEnabled(False)
                self.btn_select_panel_info.setEnabled(False)
                self.btn_del_panel.setEnabled(False)
                self.btn_update_number_ok.setVisible(True)
                self.btn_update_number_no.setVisible(True)
                self.MapImage.setDragMode(QGraphicsView.RubberBandDrag)  # 设置为可以拉框

                self.bool_updage_number = 1


        else:
            QMessageBox.information(self, "消息", "请将操作状态改为查看才可选择要重新编号的组件", QMessageBox.Yes)

    def update_number_Ok(self):
        #确认修改编号
        print("确认修改编号")
        if self.bool_updage_number == 1:
            # 修改编号
            self.update_number_dialog = Update_Number_dialog(self)
            self.update_number_dialog.Update_Number_dialog_Signal.connect(self.getDialogSignal)
            self.update_number_dialog.show()

    def update_number_No(self):
        #  取消修改编号
        print("取消修改编号")
        self.btn_update_number.setEnabled(True)
        self.btn_move_panel.setEnabled(True)
        self.btn_batch_update_panel.setEnabled(True)
        self.btn_batch_del_panel.setEnabled(True)
        self.btn_GPS_range_computing.setEnabled(True)
        self.btn_select_panel_info.setEnabled(True)
        self.btn_del_panel.setEnabled(True)
        self.btn_update_number_ok.setVisible(False)
        self.btn_update_number_no.setVisible(False)
        # 模式还原到使用之前
        if self.movemode == 1:
            self.MapImage.setDragMode(QtWidgets.QGraphicsView.ScrollHandDrag)
        else:
            self.MapImage.setDragMode(QtWidgets.QGraphicsView.NoDrag)
    #list单点事件
    def select_panel_item(self):
        print("点击事件")
        number_value = self.listPanelID.currentItem().text()
        for panel_item_info in self.panel_info_all_list:
            panel_info_dict = json.loads(panel_item_info[1])  # 转为字典

            if number_value == panel_info_dict["panel_name"]:
                for box in self.general_panel_showlist:
                    if box.key in panel_item_info[0]:
                        box.setSelected(True)
                    else:
                        box.setSelected(False)

    def select_right_menu(self,pos):
        number_Item = self.listPanelID.currentItem()
        select_panel_info_dict = []
        for panel_item_info in self.panel_info_all_list:
            panel_info_dict = json.loads(panel_item_info[1])  # 转为字典

            if number_Item.text() == panel_info_dict["panel_name"]:
                for box in self.general_panel_showlist:
                    if box.key in panel_item_info[0]:
                        select_panel_info_dict = panel_info_dict
                        box.setSelected(True)
                    else:
                        box.setSelected(False)
        print("右键点击",number_Item.text())
        menu = QMenu()
        opt1 = menu.addAction("修改信息")
        opt2 = menu.addAction("删除")
        action = menu.exec_(QtGui.QCursor.pos())
        if action == opt1:
            self.update_select_panel = Select_Update_Panel_info()
            self.update_select_panel.setT(select_panel_info_dict,number_Item,self.listPanelID)
            self.update_select_panel.Select_Update_Panel_info_Signal.connect(self.getDialogSignal)
            self.update_select_panel.show()
        if action == opt2:
            reply = QMessageBox.question(self,"删除提示","你确定要删除编号["+number_Item.text()+"]面板?",QMessageBox.Yes|QMessageBox.No)
            if reply == QMessageBox.Yes:
                print("删除当前面板")
                for panelinfo in self.panel_info_all_list:
                    panel_info_dict = json.loads(panelinfo[1])  # 转为字典
                    if number_Item.text() == panel_info_dict["panel_name"]:
                        panelId = panelinfo[0]
                        sqlitefun.sqlit_delete_panelinfo_json(self.projecttime, panelId)
                        self.panel_info_all_list.remove(panelinfo)
                self.getlSignal_refresh([0,panelId])

    def del_panel_info(self):
        if self.bool_updage_number==1:
            self.te_bottom_hint.setText("正在修改组件编号，不可删除组件")
        elif self.bool_batch_update_panel == 1:
            self.te_bottom_hint.setText("正在批量修改组件信息，不可删除组件")
        elif self.bool_region_move_panel == 1:
            self.te_bottom_hint.setText("正在移动组件，不可删除组件")
        elif self.bool_batch_del_panel ==1:
            self.te_bottom_hint.setText("正在批量删除，不可进行删除单个组件")
        else:
            #删除面板
            number_Item = self.listPanelID.currentItem()
            #判断是否选中组件
            if number_Item == None:
                QMessageBox.information(self, "消息", "您没有选择要删除的组件，请选择组件。", QMessageBox.Yes)
            else:
                reply = QMessageBox.question(self, "删除提示", "你确定要删除编号[" + number_Item.text() + "]面板?",
                                             QMessageBox.Yes | QMessageBox.No)
                if reply == QMessageBox.Yes:
                    for panelinfo in self.panel_info_all_list:
                        panel_info_dict = json.loads(panelinfo[1])  # 转为字典
                        if number_Item.text() == panel_info_dict["panel_name"]:
                            panelId = panelinfo[0]
                            sqlitefun.sqlit_delete_panelinfo_json(self.projecttime, panelId)
                            self.panel_info_all_list.remove(panelinfo)
                    self.getlSignal_refresh([0,panelId])
    def update_panel_info(self):
        if self.bool_updage_number==1:
            self.te_bottom_hint.setText("正在修改组件编号，不可修改组件信息")
        elif self.bool_batch_update_panel == 1:
            self.te_bottom_hint.setText("正在批量修改组件信息，不可修改组件信息")
        elif self.bool_region_move_panel == 1:
            self.te_bottom_hint.setText("正在移动组件，不可修改组件信息")
        elif self.bool_batch_del_panel ==1:
            self.te_bottom_hint.setText("正在批量删除，不可修改组件信息")
        else:
            #修改面板
            number_Item = self.listPanelID.currentItem()
            if number_Item == None:
                QMessageBox.information(self, "消息", "您没有选择要修改的组件，请选择组件。", QMessageBox.Yes)
            else:
                select_panel_info_dict = []
                self.panel_info_all_list = sqlitefun.sqlite_select_panelinfolist(self.projecttime)  # 获取全部组件列表信息
                for panel_item_info in self.panel_info_all_list:
                    panel_info_dict = json.loads(panel_item_info[1])  # 转为字典

                    if number_Item.text() == panel_info_dict["panel_name"]:
                        for box in self.general_panel_showlist:
                            if box.key in panel_item_info[0]:
                                select_panel_info_dict = panel_info_dict
                                box.setSelected(True)
                            else:
                                box.setSelected(False)
                self.update_select_panel = Select_Update_Panel_info()
                self.update_select_panel.setT(select_panel_info_dict, number_Item,self.listPanelID)
                self.update_select_panel.Select_Update_Panel_info_Signal.connect(self.getDialogSignal)
                self.update_select_panel.show()
    def keyPressEvent(self, event):
        # 键盘监听
        if event.key() == QtCore.Qt.Key_P:
            self.MapimageZoonOut()
            
        if event.key() == QtCore.Qt.Key_O:
            self.MapimageZoonIn()
          
        if event.key() == QtCore.Qt.Key_I:
            self.MapimageZoonRest()

        if event.key() == QtCore.Qt.Key_F:
            self.changeMoveMode()
        
        if event.key() == QtCore.Qt.Key_G:
            self.getpanel_wh()

        if event.key() == QtCore.Qt.Key_H:
            self.setpanelInfo()
        if event.key() == QtCore.Qt.Key_M:
            #设置组件移动
            self.region_loc_alter()
        if event.key() == QtCore.Qt.Key_Escape:
            #取消组件移动
            self.end_fun_no()
        if event.key() == QtCore.Qt.Key_Delete:
            self.del_panel_info()
        if event.key() == QtCore.Qt.Key_U:
            #组件修改
            self.update_panel_info()
        
        
        
    def numDialog(self,type):
        if(type == 0 ):
            win_Num_Dialog = NumDialog(self)
            win_Num_Dialog.NumDialog_Signal.connect(self.getNumvalue)
            win_Num_Dialog.show()
        if(type == 1):
            win_Row_Column_Dialog = Row_Column_Dialog(self)
            win_Row_Column_Dialog.Row_Column_Dialog_Signal.connect(self.getNumvalue)
            win_Row_Column_Dialog.show()

    def setpanelInfo(self):
        if self.bool_updage_number==1:
            self.te_bottom_hint.setText("正在修改组件编号，不能设置组件信息")
        elif self.bool_batch_update_panel == 1:
            self.te_bottom_hint.setText("正在批量修改组件信息，不可设置组件信息")
        elif self.bool_region_move_panel == 1:
            self.te_bottom_hint.setText("正在移动组件，不可设置组件信息")
        elif self.bool_batch_del_panel ==1:
            self.te_bottom_hint.setText("正在批量删除，不可设置组件信息")
        else:
            win_SetPanelInfo = SetPanelInfo(self)
            win_SetPanelInfo.SetPanelInfo_Signal.connect(self.getDialogSignal)
            win_SetPanelInfo.show()


    def getpanel_wh(self):
        print("获取组件宽高")
        if self.bool_updage_number==1:
            self.te_bottom_hint.setText("正在修改组件编号，不可获取组件宽高")
        elif self.bool_batch_update_panel == 1:
            self.te_bottom_hint.setText("正在批量修改组件信息，不可获取组件宽高")
        elif self.bool_region_move_panel == 1:
            self.te_bottom_hint.setText("正在移动组件，不可获取组件宽高")
        elif self.bool_batch_del_panel ==1:
            self.te_bottom_hint.setText("正在批量删除，不可获取组件宽高")
        else:
            self.bool_getpanel_wh = 1



    def create_project(self):
        print("创建标注项目")
        win_CreateProject = CreateProject(self)
        if self.projecttime != "":
            win_CreateProject.setProjectInfo(self.projecttime)
        else:
            win_CreateProject.setProjectInfo("")
        win_CreateProject.createproject_Signal.connect(self.getDialogSignal)
        win_CreateProject.show()
    
    def history_project(self):
        print("打开标注项目")
        win_OpenProject = OpenProject(self)
        if self.projecttime != "":
            win_OpenProject.setProjectInfo(self.projecttime)
        else:
            win_OpenProject.setProjectInfo("")
        win_OpenProject.OpenProject_Signal.connect(self.getDialogSignal)
        win_OpenProject.show()

    def export_gps_image(self):
        Folderpath = filedialog.asksaveasfilename(title=u'保存全景图', filetypes=[("TXT",".txt")],initialfile=self.projectName+".txt")
        if Folderpath:
            imagepath = Folderpath.split('.')[0]+".jpg"
            #获得gps值
            #gps_list = []

            fp = open(Folderpath, 'w')
            gps_info = ""
            for panel_item in self.panel_info_all_list:
                panel_dict = json.loads(panel_item[1])
                gps_longitude = panel_dict["panel_gps"].split(',')[0]#精度
                gps_latitude = panel_dict["panel_gps"].split(',')[1]#纬度
                pos_x = panel_dict["panel_pos"].split(',')[0]
                pos_y = panel_dict["panel_pos"].split(',')[1]
                panel_w = panel_dict["panel_wh"].split(',')[0]
                panel_h = panel_dict["panel_wh"].split(',')[1]
                panel_arrange = panel_dict["panel_arrange"]
                gps_info = panel_dict["panel_name"]+" "+gps_latitude+" "+gps_longitude+" "+panel_dict["panel_altitude"]+" "+pos_x+" "+pos_y+" "+panel_w+" "+panel_h+" "+panel_dict["panel_angle"]+" "+panel_arrange
                fp.write(gps_info)

                fp.write('\n')
            fp.close()


            print("导出gps",Folderpath)
            #导出image
            print("导出image",imagepath)
            panoramic_bin = sqlitefun.sqlite_select_imagebin(self.projecttime, "panoramic")[0][0]
            with open(imagepath, 'wb') as f:
                f.write(panoramic_bin)



    def import_gps_image(self):
        FilePath = filedialog.askopenfilename()

        if FilePath:
            if FilePath:
                imagepath = FilePath.split('.')[0] + ".jpg"
            f = open(FilePath,'r')
            data = f.readlines()
            f.close()
            info_all_list = []
            timeindex = 0
            for data_info in data:
                timeindex+=1
                data_info = data_info.strip('\n')
                panel_info = data_info.split(' ')
                panel_name_value = panel_info[0]
                panel_info_len = len(panel_info)
                panel_type_value = sqlitefun.sqlite_select_config("label_type")
                panel_model_value = sqlitefun.sqlite_select_config("panel_model")
                panel_manufactor_value= sqlitefun.sqlite_select_config("panel_manufactor")
                panel_manufactor_value = panel_manufactor_value[0][0]
                panel_type_value = panel_type_value[0][0]
                panel_model_value = panel_model_value[0][0]
                panel_power_value = sqlitefun.sqlite_select_config("panel_power")
                panel_power_value = panel_power_value[0][0]
                panel_color = sqlitefun.sqlite_select_config("panel_color")
                panel_color = panel_color[0][0]
                panel_pos_value = panel_info[4]+","+panel_info[5]
                panel_wh_value = panel_info[6]+","+panel_info[7]
                panel_gps_value = panel_info[2]+","+panel_info[1]
                panel_altitude_value = panel_info[3]
                if panel_info_len < 9:
                    panel_angle_value = "0"
                else:
                    panel_angle_value = panel_info[8]
                if len(panel_info)>9:
                    panel_arrange_value = panel_info[9]
                else:
                    panel_arrange_value = "2x10"
                t = int(time.time())
                panel_name_value = str(t)+","+panel_name_value
                panel_index_text = datetime.datetime.now().strftime('%Y%m%d_%H%M%S_%f')
                panel_index_text = panel_index_text+str(timeindex)
                self.dict_value = {'panel_index':panel_index_text,'panel_pos': panel_pos_value,
                                   'panel_name': panel_name_value, 'panel_manufactor':panel_manufactor_value,
                                   'label_type': panel_type_value, 'panel_model': panel_model_value,
                                   'panel_power': panel_power_value, 'panel_wh': panel_wh_value,
                                   'panel_altitude': panel_altitude_value, 'panel_angle': panel_angle_value,
                                   'panel_arrange':panel_arrange_value,'panel_gps': panel_gps_value,
                                   'panel_color':panel_color
                                   }
                panel_info_json = json.dumps(self.dict_value, ensure_ascii=False)
                info_all_list.append([panel_index_text, panel_info_json])
            sqlitefun.sqlite_insert_panelinfo_all_json(self.projecttime, info_all_list)
            reply = QMessageBox.question(self, "导入提示", "你是否要导入全景图?",
                                         QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.Yes:
                print("导入全景图", imagepath)
                basicsfun.savaImagebinSqlite(self.projecttime, "panoramic", imagepath)

            self.getlSignal_refresh([3])


    # def import_image(self):
    #     FilePath = filedialog.askopenfilename()
    #     print("导入全景图",FilePath)

    def getDialogSignal(self, lists):
        print("主界面收到各种信号",lists)
        #修改编号
        if lists[0] == "Update_Number_dialog":
            number_value = lists[1]
            direction_value = lists[2]
            if number_value == 0 :
                self.bool_updage_number = 0
            else:
                #获取选中点
                selectBox = []
                for box in self.general_panel_showlist:
                    if box.isSelected():
                        for panel_info in self.panel_info_all_list:
                            boxkey = box.key
                            if boxkey == panel_info[0]:
                                selectBox.append(panel_info)
                        print("boxkey", boxkey)
                # 判断选中得值中有没有此编号，如果有直可以修改，如果没有则在全部列表中寻找
                Is_select_number = False
                Is_panel_number = False
                Is_update_number = False
                for selectItem in selectBox:
                    panel_info_dict = json.loads(selectItem[1])  # 转为字典
                    panel_name = panel_info_dict["panel_name"]
                    if number_value == panel_name:
                        Is_select_number = True
                        break
                        # 如果有会替换原来得数据
                if Is_select_number == False:
                    for panel_info in self.panel_info_all_list:
                        panel_dict = json.loads(panel_info[1])
                        panel_name = panel_dict["panel_name"]
                        if number_value == panel_name:
                            Is_panel_number = True
                            break
                    if Is_panel_number == False:
                        Is_update_number = True
                    else:
                        Is_update_number = False
                else:
                    Is_update_number = True

                if Is_update_number:
                    #获取左下角坐标值
                    pos_x_list=[] #所有得x坐标
                    pos_y_list=[] #所有得y坐标

                    for row_info in selectBox:
                        panel_info_dict = json.loads(row_info[1])  # 转为字典
                        panel_pos_y = int(panel_info_dict["panel_pos"].split(',')[1])
                        panel_pos_x = int(panel_info_dict["panel_pos"].split(',')[0])
                        pos_x_list.append(panel_pos_x)
                        pos_y_list.append(panel_pos_y)
                    #左下角 x最小 y最大
                    min_x = min(pos_x_list)
                    max_y = max(pos_y_list)
                    min_y = min(pos_y_list)
                    max_x = max(pos_x_list)
                    #获取最左侧的数据
                    left_cloum_panel_list = []  # 左列
                    left_cloum_panel_list_y = [] #左列Y值
                    left_sort_panel_list = []

                    for row_info in selectBox:
                        panel_info_dict = json.loads(row_info[1])  # 转为字典
                        panel_name = panel_info_dict["panel_name"]
                        panel_pos_y = int(panel_info_dict["panel_pos"].split(',')[1])
                        panel_pos_x = int(panel_info_dict["panel_pos"].split(',')[0])
                        panel_h = int(panel_info_dict["panel_wh"].split(',')[1])
                        panel_w = int(panel_info_dict["panel_wh"].split(',')[0])
                        if abs(int(panel_pos_x-min_x)) < int(panel_w / 2):
                            #最左列
                            left_cloum_panel_list.append(panel_info_dict["panel_pos"])
                            left_cloum_panel_list_y.append(panel_pos_y)
                    left_cloum_panel_list_y.sort(reverse=True)
                    for i in range(len(left_cloum_panel_list_y)):
                        for cloum_pos in left_cloum_panel_list:
                            cloum_pos_y = cloum_pos.split(",")[1]
                            if str(left_cloum_panel_list_y[i]) ==  cloum_pos_y:
                                left_sort_panel_list.append(cloum_pos)
                                break
                    print("最左侧值排序后的值：",left_sort_panel_list)
                    Number_panel_list=[]
                    #然后找 选中 跟左侧列的y值范围不大于组件高度的值为同一行
                    for i in range(len(left_sort_panel_list)):
                        left_panel_y = int(left_sort_panel_list[i].split(",")[1])
                        row_list = []
                        sort_row_list = []
                        for selectItem in selectBox:
                            panel_info_dict = json.loads(selectItem[1])  # 转为字典
                            panel_name = panel_info_dict["panel_name"]
                            panel_pos_y = int(panel_info_dict["panel_pos"].split(',')[1])
                            panel_pos_x = int(panel_info_dict["panel_pos"].split(',')[0])
                            panel_pos_w = int(panel_info_dict["panel_wh"].split(',')[0])
                            panel_pos_h = int(panel_info_dict["panel_wh"].split(',')[1])
                            if abs(panel_pos_y - left_panel_y) < int(panel_pos_h / 2):
                                row_list.append(selectItem)
                                sort_row_list.append(panel_pos_x)
                        sort_row_list.sort()
                        print("x升序排序：", sort_row_list)
                        sort_row_panel_list = []
                        for sort_row in sort_row_list:
                            for row_panel in row_list:
                                row_panel_dict = json.loads(row_panel[1])
                                row_panel_pos_x = int(row_panel_dict["panel_pos"].split(',')[0])
                                if sort_row == row_panel_pos_x:
                                    sort_row_panel_list.append(row_panel)

                        Number_panel_list.append(sort_row_panel_list)
                    print("WQPNumber",Number_panel_list)
                    # 编号分割
                    number_sp = number_value.split(',')
                    number_format = ""
                    number_qu = number_sp[0]
                    number_row_value = int(number_sp[1])
                    number_cloum_value = int(number_sp[2])
                    for row_panel_list in Number_panel_list:
                        #number_row_value = int(number_sp[1])
                        number_cloum_value = int(number_sp[2])
                        #每一行数据
                        for row_panel_item in row_panel_list:
                            new_number_value = number_qu + "," + str(number_row_value) + "," + str(number_cloum_value)
                            row_panel_dict = json.loads(row_panel_item[1])
                            row_panel_dict["panel_name"] = new_number_value
                            row_panel_json = json.dumps(row_panel_dict, ensure_ascii=False)
                            sqlitefun.sqlite_update_panelinfo_json(self.projecttime, row_panel_item[0],
                                                                   row_panel_json)
                            if "右增" in direction_value:
                                number_cloum_value += 1
                            if "右减" in direction_value:
                                number_cloum_value -= 1
                        if "上增" in direction_value:
                            number_row_value += 1
                        if "上减" in direction_value:
                            number_row_value -= 1

                    self.getlSignal_refresh([2])
                    self.update_number_dialog.close()
                    #修改完成后重置状态
                    self.bool_updage_number = 0 #修改重新编号状态
                    self.btn_update_number.setEnabled(True)
                    self.btn_move_panel.setEnabled(True)
                    self.btn_batch_update_panel.setEnabled(True)
                    self.btn_batch_del_panel.setEnabled(True)
                    self.btn_GPS_range_computing.setEnabled(True)
                    self.btn_select_panel_info.setEnabled(True)
                    self.btn_del_panel.setEnabled(True)
                    self.btn_update_number_no.setVisible(False)
                    self.btn_update_number_ok.setVisible(False)
                    if self.movemode == 1:
                        self.MapImage.setDragMode(QtWidgets.QGraphicsView.ScrollHandDrag)
                    else:
                        self.MapImage.setDragMode(QtWidgets.QGraphicsView.NoDrag)
                    self.te_bottom_hint.setText("编号修改完成。")
                else:
                    QMessageBox.information(self, "消息", "此命名已存在，请修改初始编号", QMessageBox.Yes)


        #修改面板信息
        elif lists[0] == "Select_Update_Panel_info":
            #修改内容
            dict_value = lists[1]
            listItem = lists[2]
            batch_update_list = lists[3]
            isUpdate = lists[4]
            update_info_list = []
            if isUpdate == 0:
                self.update_select_panel.close()
                self.bool_batch_update_panel = 0
            else:
                if listItem != None:
                    update_info_list.clear()
                    selectIndex = self.listPanelID.row(listItem)
                    select_panel_info = self.panel_info_all_list[selectIndex]
                    panel_info_dict = json.loads(select_panel_info[1])  # 转为字典
                    panel_info_dict["panel_name"] = dict_value['panel_name']
                    panel_info_dict["panel_manufactor"] = dict_value['panel_manufactor']
                    panel_info_dict["label_type"] = dict_value['label_type']
                    panel_info_dict["panel_model"] = dict_value['panel_model']
                    panel_info_dict["panel_power"] =dict_value['panel_power']
                    panel_info_dict["panel_wh"] =  dict_value['panel_wh']
                    panel_info_dict["panel_altitude"] = dict_value['panel_altitude']
                    panel_info_dict["panel_angle"] = dict_value['panel_angle']
                    panel_info_dict["panel_arrange"] = dict_value['panel_arrange']
                    panel_info_dict["panel_gps"] = dict_value['panel_gps']
                    panel_info_dict["panel_color"] = dict_value['panel_color']

                    panel_info_json = json.dumps(panel_info_dict, ensure_ascii=False)

                    sqlitefun.sqlite_update_panelinfo_json(self.projecttime, panel_info_dict["panel_index"], panel_info_json)
                    #listItem.setText(dict_value['panel_name'])
                    update_info_list.append(select_panel_info[0]+"|"+panel_info_dict["panel_color"])
                    self.update_select_panel.close()

                if len(batch_update_list)>0:
                    update_info_list.clear()
                    sqlite_panel_info_all = []
                    for select_key in batch_update_list:
                        print("WQP   ", select_key)
                        for panel_info in self.panel_info_all_list:
                            panel_info_dic = json.loads(panel_info[1])
                            if select_key == panel_info[0]:
                                #获得到需要修改的组件
                                panel_info_dic["panel_model"] = dict_value["panel_model"]
                                panel_info_dic["panel_color"] = dict_value["panel_color"]
                                panel_info_dic["panel_manufactor"] = dict_value['panel_manufactor']
                                panel_info_dic["label_type"] = dict_value['label_type']
                                panel_info_dic["panel_power"] = dict_value['panel_power']
                                panel_info_dic["panel_wh"] = dict_value['panel_wh']
                                panel_info_dic["panel_altitude"] = dict_value['panel_altitude']
                                panel_info_dic["panel_angle"] = dict_value['panel_angle']
                                panel_info_dic["panel_arrange"] = dict_value['panel_arrange']

                                update_info_list.append(select_key+"|"+panel_info_dic["panel_color"])
                                # panel_info_dic["panel_power"] = dict_value["panel_power"]
                                # panel_info_dic["panel_arrange"] = dict_value["panel_arrange"]
                                panel_info_json = json.dumps(panel_info_dic, ensure_ascii=False)
                                sqlite_panel_info_all.append([select_key, panel_info_json])
                    sqlitefun.sqlite_update_panelinfo_all_json(self.projecttime, sqlite_panel_info_all)
                    self.update_select_panel.close()
                self.getlSignal_refresh([2,update_info_list])
                self.bool_batch_update_panel = 0
        elif lists[0] == "SetPanel_Info":
            template_list = sqlitefun.sqlite_select_template_all()

            if len(template_list) > 0:
                item_index = 0
                for templat in template_list:
                    self.cb_panel_mode.addItem(templat[0])
                for count in range(self.cb_panel_mode.count()):
                    old_model = sqlitefun.sqlite_select_config("panel_model")

                    if old_model[0][0] == self.cb_panel_mode.itemText(count):
                        print("SetPanelInfo",old_model[0][0])
                        item_index = count
                        break
                self.cb_panel_mode.setCurrentIndex(item_index)
                self.getlSignal_refresh([2])
        elif lists[0] == "Update_Panel_Model_Signal":
            print("修改完信息之后的处理")
            #传过来 修改的模板名称
            panel_model = lists[1]
            model_json = lists[2]
            type = lists[3]
            if type ==1:
                model_dict = json.loads(model_json)
                # dict_value = {'panel_manufactor': panel_manufactor,
                #               'label_type': panel_type, 'panel_model': panel_model,
                #               'panel_power': panel_power, 'panel_wh': panel_wh,
                #               'panel_altitude': panel_altitude,
                #               'panel_angle': panel_angle, 'panel_arrange': panel_arrange,
                #               'panel_color': panel_color
                #               }
                # 获取数据中所有用这个模板的信息
                for panel_info in self.panel_info_all_list:
                    panel_dict = json.loads(panel_info[1])
                    if panel_dict["panel_model"] == panel_model:
                        panel_dict["panel_manufactor"] = model_dict["panel_manufactor"]
                        panel_dict["label_type"] = model_dict["label_type"]
                        panel_dict["panel_power"] = model_dict["panel_power"]
                        panel_dict["panel_wh"] = model_dict["panel_wh"]
                        panel_dict["panel_altitude"] = model_dict["panel_altitude"]
                        panel_dict["panel_angle"] = model_dict["panel_angle"]
                        panel_dict["panel_arrange"] = model_dict["panel_arrange"]
                        panel_dict["panel_color"] = model_dict["panel_color"]
                        panel_info_json = json.dumps(panel_dict, ensure_ascii=False)
                        sqlitefun.sqlite_update_panelinfo_json(self.projecttime, panel_dict["panel_index"],
                                                               panel_info_json)
                        for box in self.general_panel_showlist:
                            if box.key == panel_dict["panel_index"]:
                                box.color = self.Hex_to_RGB(panel_dict["panel_color"])

                self.getlSignal_refresh([2])
                # 修改他们的数据并保存到数据库中

                self.Update_Model.close()
            elif type ==2:
                #更新选择面板信息
                template_list = sqlitefun.sqlite_select_template_all()
                #判断当前选择的内容为
                old_model = sqlitefun.sqlite_select_config("panel_model")
                select_model =""
                if old_model[0][0] == panel_model:
                    select_model = template_list[0][0]
                    panel_json = sqlitefun.sqlite_select_template(select_model)
                    panel_dict = json.loads(panel_json[0][0])
                    sqlitefun.sqlite_insert_config("label_type", panel_dict["label_type"])
                    sqlitefun.sqlite_insert_config("panel_model", panel_dict["panel_model"])
                    sqlitefun.sqlite_insert_config("panel_power", panel_dict["panel_power"])
                    sqlitefun.sqlite_insert_config("panel_wh", panel_dict["panel_wh"])
                    sqlitefun.sqlite_insert_config("panel_altitude", panel_dict["panel_altitude"])
                    sqlitefun.sqlite_insert_config("panel_angle", panel_dict["panel_angle"])
                    sqlitefun.sqlite_insert_config("panel_manufactor", panel_dict["panel_manufactor"])
                    sqlitefun.sqlite_insert_config("panel_arrange", panel_dict["panel_arrange"])
                    sqlitefun.sqlite_insert_config("panel_color", panel_dict["panel_color"])
                    mode_info = select_model + "模型内容信息为：\n"
                    mode_info += "标注类型：" + panel_dict["label_type"] + ","
                    mode_info += "组件型号：" + panel_dict["panel_model"] + ","
                    mode_info += "组件功率：" + panel_dict["panel_power"] + ","
                    mode_info += "组件宽高：" + panel_dict["panel_wh"] + ","
                    mode_info += "组件海拔：" + panel_dict["panel_altitude"] + ","
                    mode_info += "组件角度：" + panel_dict["panel_angle"] + ","
                    mode_info += "组件厂家：" + panel_dict["panel_manufactor"] + ","
                    mode_info += "组件排布：" + panel_dict["panel_arrange"]
                else:
                    select_model = old_model[0][0]


                self.cb_panel_mode.clear()
                print("ssss")
                if len(template_list) > 0:
                    item_index = 0
                    for templat in template_list:
                        self.cb_panel_mode.addItem(templat[0])
                    for count in range(self.cb_panel_mode.count()):
                        #old_model = sqlitefun.sqlite_select_config("panel_model")
                        if select_model == self.cb_panel_mode.itemText(count):
                            item_index = count
                            break
                    self.cb_panel_mode.setCurrentIndex(item_index)



        else:
            self.projecttime = lists[1]  # 获取当前打开项目的索引值
            self.projectName = lists[2]
            old_project_time = lists[3] #上次打开的projectTime
            if old_project_time != "":
                print("")
                self.get_old_template(old_project_time)
            self.lab_projectName.setText(lists[1]+" / "+lists[2])
            self.cb_label_mode.setCurrentText("查看")
            sqlitefun.sqlite_create_panelinfo(self.projecttime)
            print("调试 看一下字典信息对不对",self.panel_info_dictlist)
            if lists[0] == "CreateProject":
                panoramic_bin = sqlitefun.sqlite_select_imagebin(lists[1], "panoramic")[0][0]
                self.showPanoramic(panoramic_bin)

            if lists[0] == "OpenProject":
                panoramic_bin = sqlitefun.sqlite_select_imagebin(lists[1],"panoramic")[0][0]
                self.showPanoramic(panoramic_bin)
            self.showPanelList()
    #获取前项目 使用的模板数据
    def get_old_template(self,projectTime):
        old_template_list = []
        sqlitefun.sql_create_use_template()
        old_panel_info_list =  sqlitefun.sqlite_select_panelinfolist(projectTime)
        for old_panel_info in old_panel_info_list:
            old_panel_dict = json.loads(old_panel_info[1])
            old_temple = old_panel_dict["panel_model"]
            old_template_list.append(old_temple)
        old_template_list = set(old_template_list)
        print("上一个项目使用过的模板",old_template_list)
        #查询所有已使用模板
        all_use_model = sqlitefun.sqlite_select_all_use_template()
        old_model_list = []
        delete_model_list = []

        for use_model in all_use_model:
            if projectTime in use_model[1]:
                old_model_list.append(use_model[0])
        print("old_model_list",old_model_list)
        for old_model_name in old_model_list:
            if old_model_name not in old_template_list:
                delete_model_list.append(old_model_name)
        print("要删除的组件信息",delete_model_list)
        update_use_projectid = ""
        #删除组件信息
        for delete_model in delete_model_list:
            result = sqlitefun.sqlite_select_use_template(delete_model)
            projectid_list = result[0][0].split(';')
            for projectid in projectid_list:
                if projectid != projectTime and projectid != '':
                    update_use_projectid += projectid+";"
            sqlitefun.sqlite_update_use_template(delete_model, update_use_projectid)

        print("更新后的id",update_use_projectid)

        for old_temple in old_template_list:
            result = sqlitefun.sqlite_select_use_template(old_temple)
            if len(result) > 0:
                if projectTime not in result[0][0]:
                   project_id = result[0][0]+projectTime+";"
                   sqlitefun.sqlite_update_use_template(old_temple,project_id)
            else:
                self.insert_panel_model(old_temple, projectTime+";")


    def getNumvalue(self,list):

        if list[0] == "NumDialog": #单行标注
            self.numValue = int(list[1])
            if self.numValue == -1:
                #取消按钮
                if len(self.select_pos) > 0:
                    self.delete_pos(self.select_pos)
                    self.select_pos = []
                    self.rownum = 2
            else:
                #确认按钮事件
                rowvalue1_x = int(self.rowvalue1[0])
                rowvalue2_x = int(self.rowvalue2[0])
                rowvalue1_y = int(self.rowvalue1[1])
                rowvalue2_y = int(self.rowvalue2[1])

                value_x = (abs(rowvalue2_x-rowvalue1_x))/(self.numValue-1)
                value_y = (rowvalue2_y-rowvalue1_y)/(self.numValue-1)
                num= int(list[1]) -2
                index = 0
                panel_pos_int_w = int(self.panel_wh_text.split(",")[0])
                panel_pos_int_h = int(self.panel_wh_text.split(",")[1])
                while index < num:
                    index+=1
                    if(abs(int(value_x)>5) and abs(int(value_y)<5)):
                        #横向
                        if(rowvalue1_x < rowvalue2_x):
                            map_x = (rowvalue1_x + (value_x* index));
                            map_y = rowvalue1_y + (value_y*index)
                        else:
                            map_x = (rowvalue2_x + (value_x  * index));
                            map_y = rowvalue2_y - (value_y*index)
                    if(abs(int(value_x))<5 and abs(int(value_y))>5):
                        #竖向
                        if (rowvalue1_y < rowvalue2_y):
                            map_x = rowvalue1_x + (value_x * index)
                            map_y = (rowvalue1_y + (value_y * index));
                        else:
                            map_x = rowvalue2_x - (value_x * index)
                            map_y = (rowvalue2_y - (value_y * index));
                    if(abs(int(value_x)>5) and abs(int(value_y)>5)):
                        #斜
                        if(rowvalue1_x < rowvalue2_x) and (rowvalue1_y< rowvalue2_y):
                            map_x = (rowvalue1_x + (value_x * index));
                            map_y = (rowvalue1_y + (value_y * index));
                            print("1")
                        elif (rowvalue1_x > rowvalue2_x) and (rowvalue1_y < rowvalue2_y):
                            map_x = (rowvalue1_x - (value_x * index));
                            map_y = (rowvalue1_y + (value_y * index));
                            print("2")
                        elif (rowvalue1_x < rowvalue2_x) and (rowvalue1_y > rowvalue2_y):
                            map_x = (rowvalue2_x - (value_x* index));
                            map_y = (rowvalue2_y + (value_y * index));
                            print("3")
                        elif (rowvalue1_x > rowvalue2_x) and (rowvalue1_y > rowvalue2_y):
                            map_x = (rowvalue2_x + (value_x * index));
                            map_y = (rowvalue2_y + (value_y * index));
                            print("4")

                    if (abs(int(value_x) < 5) and abs(int(value_y) < 5)):
                        # 斜
                        if (rowvalue1_x < rowvalue2_x) and (rowvalue1_y < rowvalue2_y):
                            map_x = (rowvalue1_x - (value_x * index));
                            map_y = (rowvalue1_y + (value_y * index));
                            print("1")
                        elif (rowvalue1_x > rowvalue2_x) and (rowvalue1_y < rowvalue2_y):
                            map_x = (rowvalue1_x + (value_x * index));
                            map_y = (rowvalue1_y + (value_y * index));
                            print("2")
                        elif (rowvalue1_x < rowvalue2_x) and (rowvalue1_y > rowvalue2_y):
                            map_x = (rowvalue2_x + (value_x * index));
                            map_y = (rowvalue2_y + (value_y * index));
                            print("3")
                        elif (rowvalue1_x > rowvalue2_x) and (rowvalue1_y > rowvalue2_y):
                            map_x = (rowvalue2_x - (value_x * index));
                            map_y = (rowvalue2_y + (value_y * index));
                            print("4")

                    arg1 = [map_x, map_y]
                    self.add_panel_info(arg1)
                self.select_pos=[]
            self.te_bottom_hint.setText("单行标注完成")
        if list[0] == "Row_Column_Dialog": #多行标注
            Columnvalue = int(list[1])
            rowValue = int(list[2])
            if Columnvalue == -1 and rowValue == -1:
                if len(self.select_pos) > 0:
                    self.delete_pos(self.select_pos)
                    self.select_pos = []
                    self.row_column = 4
            else:
                row_column1 = self.rowcolumn[0]
                row_column2 = self.rowcolumn[1]
                row_column3 = self.rowcolumn[2]
                row_column4 = self.rowcolumn[3]
                value1_x = (abs(row_column4[0] - row_column1[0])) / (Columnvalue - 1)
                value1_y = (row_column4[1] - row_column1[1]) / (Columnvalue - 1)
                value2_x = (abs(row_column3[0] - row_column2[0])) / (Columnvalue - 1)
                value2_y = (row_column3[1] - row_column2[1]) / (Columnvalue - 1)
                column_num = Columnvalue - 2 #列
                row_num = rowValue -2 #行
                column_arr1 = []
                column_arr2 = []
                panel_pos_int_w = int(self.panel_wh_text.split(",")[0])
                column_arr1.append(row_column1)
                column_arr1.append(row_column4)
                column_arr2.append(row_column2)
                column_arr2.append(row_column3)
                index = 0
                while index < column_num:
                    index+=1
                    if (abs(int(value1_x)) < 5 and abs(int(value1_y)) > 5):
                        #竖向
                        if (row_column1[1] < row_column4[1]):
                            map_x_1 = row_column1[0] + (value1_x * index)
                            map_y_1 = row_column1[1] + (value1_y * index)
                        else:
                            map_x_1 = row_column4[0] + (value1_x * index)
                            map_y_1 = row_column4[1] + (value1_y * index)
                    elif (abs(int(value1_x)) > 5 and abs(int(value1_y)) > 5):
                        #斜方向
                        # 斜方向
                        coordinate1_x = row_column1[0]
                        coordinate2_x = row_column4[0]
                        coordinate1_y = row_column1[1]
                        coordinate2_y = row_column4[1]
                        if (coordinate1_x < coordinate2_x) and (coordinate1_y < coordinate2_y):
                            map_x_1 = (coordinate1_x + (value1_x * index));
                            map_y_1 = (coordinate1_y + (value1_y * index));
                            print("1")
                        elif (coordinate1_x > coordinate2_x) and (coordinate1_y < coordinate2_y):
                            map_x_1 = (coordinate1_x - (value1_x * index));
                            map_y_1 = (coordinate1_y + (value1_y * index));
                            print("2")
                        elif (coordinate1_x < coordinate2_x) and (coordinate1_y > coordinate2_y):
                            map_x_1 = (coordinate2_x - (value1_x * index));
                            map_y_1 = (coordinate2_y + (value1_y * index));
                            print("3")
                        elif (coordinate1_x > coordinate2_x) and (coordinate1_y > coordinate2_y):
                            map_x_1 = (coordinate2_x + (value1_x * index));
                            map_y_1 = (coordinate2_y + (value1_y * index));
                            print("4")
                    if (abs(int(value2_x)) < 5 and abs(int(value2_y)) > 5):
                        if (row_column2[1] < row_column3[1]):
                            map_x_2 = row_column2[0] + (value2_x * index)
                            map_y_2 = row_column2[1] + (value2_y * index)
                        else:
                            map_x_2 = row_column3[0] + (value2_x * index)
                            map_y_2 = row_column3[1] + (value2_y * index)
                    elif (abs(int(value2_x)) > 5 and abs(int(value2_y)) > 5):
                        #斜方向
                        coordinate1_x = row_column2[0]
                        coordinate2_x = row_column3[0]
                        coordinate1_y = row_column2[1]
                        coordinate2_y = row_column3[1]
                        if (coordinate1_x < coordinate2_x) and (coordinate1_y < coordinate2_y):
                            map_x_2 = (coordinate1_x + (value2_x * index));
                            map_y_2 = (coordinate1_y + (value2_y * index));
                            print("1")
                        elif (coordinate1_x > coordinate2_x) and (coordinate1_y < coordinate2_y):
                            map_x_2 = (coordinate1_x - (value2_x * index));
                            map_y_2 = (coordinate1_y + (value2_y * index));
                            print("2")
                        elif (coordinate1_x < coordinate2_x) and (coordinate1_y > coordinate2_y):
                            map_x_2 = (coordinate2_x - (value2_x * index));
                            map_y_2 = (coordinate2_y + (value2_y * index));
                            print("3")
                        elif (coordinate1_x > coordinate2_x) and (coordinate1_y > coordinate2_y):
                            map_x_2 = (coordinate2_x + (value2_x * index));
                            map_y_2 = (coordinate2_y + (value2_y * index));
                            print("4")

                    arg1 = [map_x_1, map_y_1]
                    arg2 = [map_x_2, map_y_2]
                    column_arr1.append(arg1)
                    column_arr2.append(arg2)
                    self.add_panel_info(arg1)
                    self.add_panel_info(arg2)
                for i in range(Columnvalue):
                    value1_x = column_arr1[i][0]
                    value1_y = column_arr1[i][1]
                    value2_x = column_arr2[i][0]
                    value2_y = column_arr2[i][1]
                    value_x = (abs(value2_x - value1_x)) / (rowValue - 1)
                    value_y = (value2_y - value1_y) / (rowValue - 1)
                    row_index = 0
                    while row_index<row_num:
                        row_index += 1
                        if (int(value_x) > 5 and int(value_y) < 5):
                            # 横向
                            if (value1_x < value2_x):
                                map_x = value1_x + (value_x  * row_index);
                                map_y = value1_y + (value_y * row_index)
                            else:
                                map_x = value2_x + (value_x * row_index);
                                map_y = value2_y + (value_y * row_index)
                        arg1 = [map_x, map_y]
                        self.add_panel_info(arg1)
                self.select_pos = []
            self.te_bottom_hint.setText("多行标注完成")


    def showPanoramic(self,image_bin):
        # 显示全景图地图
        print("显示全景地图")
        self.general_panel_showlist=[]
        self.panel_info_all_list = []
        self.panel_info_dictlist = []
        self.showMapImage()  # 初始化全景图控件
        self.loadImage(image_bin)
        self._scene.signalMousePos.connect(self.cliskMapMovepos) # 连接信号
        self.whlabel = ImagerectboxItem_wh("whlabel")
        self.menu77.setEnabled(True)
        # 将所有的按钮设置为可点击状态
        self.but_zoomout.setEnabled(True)
        self.but_movemode.setEnabled(True)
        self.but_zoomin.setEnabled(True)
        self.but_rest.setEnabled(True)
        self.but_setpanel_info.setEnabled(True)
        self.but_update_panel_mode.setEnabled(True)
        self.but_delete_panel_mode.setEnabled(True)
        self.but_getpanel_wh.setEnabled(True)
        self.btn_move_panel.setEnabled(True)
        self.btn_update_number.setEnabled(True)
        self.btn_GPS_computing.setEnabled(True)
        self.btn_GPS_range_computing.setEnabled(True)
        self.btn_GPS_undone_computing.setEnabled(True)
        self.btn_GPS_clear.setEnabled(True)
        self.btn_del_panel.setEnabled(True)
        self.btn_select_panel_info.setEnabled(True)
        self.btn_batch_update_panel.setEnabled(True)
        self.btn_batch_del_panel.setEnabled(True)
        self.cb_label_mode.setEnabled(True)
        self.cb_panel_mode.setEnabled(True)
        template_list = sqlitefun.sqlite_select_template_all()

        if len(template_list) > 0:
            item_index = 0
            for templat in template_list:
                self.cb_panel_mode.addItem(templat[0])
            for count in range(self.cb_panel_mode.count()):
                old_model = sqlitefun.sqlite_select_config("panel_model")
                if old_model[0][0] == self.cb_panel_mode.itemText(count):
                    item_index = count
                    break
            self.cb_panel_mode.setCurrentIndex(item_index)
            self.cb_panel_mode.currentIndexChanged.connect(self.templateClick)
        # 显示 基本组件图例
        self.listPanelID.clear()
        self.panel_info_dictlist.clear()
        self.panel_info_all_list = sqlitefun.sqlite_select_panelinfolist(self.projecttime)  # 获取全部组件列表信息
        for panel_info_json in self.panel_info_all_list:
            panel_info_dict = json.loads(panel_info_json[1])  # 转为字典
            self.panel_info_dictlist.append(panel_info_dict)  # 加入字典列表方便使用

            # 列表显示增加
            self.listPanelID.addItem(panel_info_dict["panel_name"])

           # 大地图显示组件增加
            panel_pos_int_w = int(panel_info_dict["panel_wh"].split(",")[0])
            panel_pos_int_h = int(panel_info_dict["panel_wh"].split(",")[1])
            panel_pos_int_x = int(float(panel_info_dict["panel_pos"].split(",")[0])) - panel_pos_int_w/2
            panel_pos_int_y = int(float(panel_info_dict["panel_pos"].split(",")[1])) - panel_pos_int_h/2

            lightGreenboxIte=PanelrectboxItem(key = panel_info_dict["panel_index"])
            lightGreenboxIte.rectpos = [panel_pos_int_x,panel_pos_int_y,panel_pos_int_w,panel_pos_int_h]
            panel_pos_model = panel_info_dict["panel_model"]
            panel_json = sqlitefun.sqlite_select_template(panel_pos_model)
            panel_dict = json.loads(panel_json[0][0])
            lightGreenboxIte.color = self.Hex_to_RGB(panel_dict["panel_color"])

            lightGreenboxIte.setFlag(QGraphicsItem.ItemIsSelectable)
            self._scene.addItem(lightGreenboxIte)
            self.general_panel_showlist.append(lightGreenboxIte)




    def cliskMapMovepos(self,arg1,arg2,mode,isright):
        # 监听大地图点击操作 包括 点击 移动 释放
        # arg1 是按下时的点
        # arg2 是移动后的点,释放时也是这个点
        # mode 1 表示 按下  0 表示 释放
        #isright 0 表示没有按下右键，1表示按下右键
        if self.movemode == 0:  # 固定状态下才允许进行其他操作
            #print("当前点击状态",self.cb_label_mode.currentText())
            print("点击参数",arg1,arg2,mode,isright)
            self.selectItem = arg1
            #self.select_pos = []


            if self.cb_label_mode.currentText() == "查看":
                #print("当前状态下允许点击组件 查看信息 和修改信息 和获取组件大小")
                if mode == 1 or mode == 0:
                    if self.bool_getpanel_wh == 1 :
                        self._scene.removeItem(self.whlabel)
                        self.whlabel.rectpos = [int(float(arg1[0])),int(float(arg1[1])),int(float(arg2[0]-arg1[0])),int(float(arg2[1]-arg1[1]))]
                        self.whlabel.color = QtGui.QPen(Qt.green,2,Qt.SolidLine)

                        self._scene.addItem(self.whlabel)
                        panel_w = str(abs(int(float(arg2[0]-arg1[0]))))
                        panel_h = str(abs(int(float(arg2[1]-arg1[1]))))
                        self.te_bottom_hint.setText("当前获取到宽高:"+"宽 "+panel_w+" 高 "+panel_h)
                        if mode == 0:
                            self.te_bottom_hint.setText("当前获取到宽高:"+"宽 "+panel_w+" 高 "+panel_h+",获取完成")
                            sqlitefun.sqlite_create_config()
                            sqlitefun.sqlite_insert_config("panel_wh",panel_w+","+panel_h)  # 写进数据库  读取好用
                            print("宽",panel_w,"高",panel_h)
                            self.bool_getpanel_wh = 0 # 左键释放 停止绘图
                if mode == 3:
                    if len(arg1)>0:
                        select_key = arg1[0].key
                        print("选中值key",select_key)

                        for panel_info in self.panel_info_all_list:
                            if select_key in panel_info[0]:
                                panel_info_dict = json.loads(panel_info[1])  # 转为字典
                                print("选中值dic", panel_info_dict)
                                listcount = self.listPanelID.count()
                                for i in range(listcount):
                                    if(panel_info_dict["panel_name"] == self.listPanelID.item(i).text()):
                                        print("select_panel_name ",self.listPanelID.item(i).text())
                                        self.listPanelID.setCurrentItem(self.listPanelID.item(i))



            if self.cb_label_mode.currentText() == "单点标注":
                if mode == 0:
                    print("当前模式 单点标注 坐标",arg1)
                    self.add_panel_info(arg1)

            if self.cb_label_mode.currentText() == "单行标注":
                #先点击两个点记录下来，选择行内个数
                if (isright == 1):
                    if self.rownum == 1:
                        self.delete_pos(self.select_pos)
                        self.select_pos=[]
                        self.rownum=2
                if mode == 0:
                    if(self.rownum==2):
                        print("debug",self.rownum)
                        self.rownum = self.rownum-1
                        self.rowvalue1=arg1
                        select_pos = str(int(self.rowvalue1[0])) + "," + str(int(self.rowvalue1[1]))
                        self.select_pos.append(select_pos)
                        self.add_panel_info(arg1)
                        self.te_bottom_hint.setText("请按照水平方向输入最后一个点")
                    elif(self.rownum==1):
                        print("debugaaaa", self.rownum)
                        self.rownum=2
                        #弹出弹出框，提示每行显示几个
                        self.rowvalue2 = arg1
                        self.add_panel_info(arg1)
                        select_pos = str(int(arg1[0])) + "," + str(int(arg1[1]))
                        self.select_pos.append(select_pos)
                        self.numDialog(0)

                pass
            if self.cb_label_mode.currentText() == "多行标注":
                if (isright == 1):
                    #self.select_pos=[]
                    if self.row_column < 4:
                        # for r_c in self.rowcolumn:
                        #     select_pos = str(int(r_c[0])) + "," + str(int(r_c[1]))
                        #     self.select_pos.append(select_pos)
                        #     print("select_pos", select_pos)
                        self.delete_pos(self.select_pos)
                        self.select_pos = []
                        self.row_column=4
                if mode == 0:
                    if (self.row_column == 1):
                        self.rowcolumn[4 - self.row_column] = arg1
                        self.add_panel_info(arg1)
                        select_pos = str(int(arg1[0])) + "," + str(int(arg1[1]))
                        self.select_pos.append(select_pos)
                        self.numDialog(1)
                        self.row_column=4
                    elif (self.row_column <= 4):
                        self.rowcolumn[4-self.row_column]= arg1
                        self.add_panel_info(arg1)
                        self.row_column = self.row_column - 1
                        select_pos = str(int(arg1[0])) + "," + str(int(arg1[1]))
                        self.select_pos.append(select_pos)
                    self.te_bottom_hint.setText("请按照顺时针方向输入第"+str(5-self.row_column)+"个点")
                pass
    def delete_pos(self,pos):
        selectBox =[]
        for pos in pos:
            for panel_info in self.panel_info_all_list:
                panel_dict = json.loads(panel_info[1])
                if pos == panel_dict["panel_pos"]:
                    key = panel_info[0]
                    for box in self.general_panel_showlist:
                        if box.key == key:
                            sqlitefun.sqlit_delete_panelinfo_json(self.projecttime, key)
                            selectBox.append(key)
                            print("单行标注删除点", key)
        self.getlSignal_refresh([4, selectBox])
    def get_panel_info(self):
        print("从数据库中获取一次当前面板的信息")

        sqlitefun.sqlite_create_panelinfo(self.projecttime)  # 创建组件信息记录表
        self.label_type_text = sqlitefun.sqlite_select_config("label_type")
        self.panel_model_text = sqlitefun.sqlite_select_config("panel_model")
        self.panel_power_text = sqlitefun.sqlite_select_config("panel_power")
        self.panel_wh_text = sqlitefun.sqlite_select_config("panel_wh")
        self.panel_altitude_text = sqlitefun.sqlite_select_config("panel_altitude")
        self.panel_angle_text = sqlitefun.sqlite_select_config("panel_angle")
        self.panel_manufactor_text = sqlitefun.sqlite_select_config("panel_manufactor")
        self.panel_arrange_text = sqlitefun.sqlite_select_config("panel_arrange")
        self.panel_color_text = sqlitefun.sqlite_select_config("panel_color")
        self.panel_color_text = self.panel_color_text[0][0]

        if len(self.label_type_text) == 0:
            self.label_type_text = "光伏组件"
        else:
            self.label_type_text = self.label_type_text[0][0]

        if len(self.panel_model_text) == 0:
            self.panel_model_text = "标准型号"
        else:
            self.panel_model_text = self.panel_model_text[0][0]

        if len(self.panel_power_text) == 0:
            self.panel_power_text = "0"
        else:
            self.panel_power_text = self.panel_power_text[0][0]
       
        if len(self.panel_wh_text) == 0:
            self.panel_wh_text = "10,10"
        else:
            self.panel_wh_text = self.panel_wh_text[0][0]

        if len(self.panel_altitude_text) == 0:
            self.panel_altitude_text = "0"
        else:
            self.panel_altitude_text = self.panel_altitude_text[0][0]

        if len(self.panel_angle_text) == 0:
            self.panel_angle_text = "0"
        else:
            self.panel_angle_text = self.panel_angle_text[0][0]

        if len(self.panel_manufactor_text) == 0:
            self.panel_manufactor_text = ""
        else:
            self.panel_manufactor_text = self.panel_manufactor_text[0][0]

        if len(self.panel_arrange_text) == 0:
            self.panel_arrange_text = "2x10"
        else:
            self.panel_arrange_text = self.panel_arrange_text[0][0]
    #将已使用的模板添加到数据库中
    def insert_panel_model(self,model_name,project_time):
        #判断表是否存在，如果不存在则创建

        sqlitefun.sqlite_insert_use_template(model_name,project_time)
    def add_panel_info(self,pos):
        print("增加面板点信息")
        self.get_panel_info() # 获取当前组件的信息
        self.panel_pos_text = str(int(pos[0]))+","+str(int(pos[1]) ) # 组件的中心坐标
        self.panel_index_text = datetime.datetime.now().strftime('%Y%m%d_%H%M%S_%f')  # 组件索引 不进行修改的
        self.panel_info_all_list = sqlitefun.sqlite_select_panelinfolist(self.projecttime)  # 获取全部组件列表信息
        self.panel_name_text ="0,"+str(len(self.panel_info_all_list)) #设置一下组件的名称 可以修改

        panel_info_dict = {}
        panel_info_dict["panel_index"] = self.panel_index_text
        panel_info_dict["panel_pos"] = self.panel_pos_text
        panel_info_dict["panel_name"] = self.panel_name_text
        panel_info_dict["label_type"] = self.label_type_text
        panel_info_dict["panel_model"] = self.panel_model_text
        panel_info_dict["panel_power"] = self.panel_power_text
        panel_info_dict["panel_wh"] = self.panel_wh_text
        panel_info_dict["panel_altitude"] = self.panel_altitude_text
        panel_info_dict["panel_angle"] = self.panel_angle_text
        panel_info_dict["panel_gps"] = "1,1"
        panel_info_dict["panel_manufactor"] = self.panel_manufactor_text
        panel_info_dict["panel_arrange"] = self.panel_arrange_text
        panel_info_dict["panel_color"] = self.panel_color_text
        panel_info_json = json.dumps(panel_info_dict,ensure_ascii=False)
        print("调试json信息",panel_info_json)
        sqlitefun.sqlite_insert_panelinfo_json(self.projecttime,self.panel_index_text,panel_info_json) # 写入或者更新组件信息


         # 列表显示增加
        self.listPanelID.addItem(panel_info_dict["panel_name"])

           # 大地图显示组件增加
        panel_pos_int_w = int(panel_info_dict["panel_wh"].split(",")[0])
        panel_pos_int_h = int(panel_info_dict["panel_wh"].split(",")[1])
        panel_pos_int_x = int(panel_info_dict["panel_pos"].split(",")[0]) - panel_pos_int_w/2
        panel_pos_int_y = int(panel_info_dict["panel_pos"].split(",")[1]) - panel_pos_int_h/2

        lightGreenboxIte=PanelrectboxItem(key = panel_info_dict["panel_index"])
        lightGreenboxIte.rectpos = [panel_pos_int_x,panel_pos_int_y,panel_pos_int_w,panel_pos_int_h]
        lightGreenboxIte.color = self.Hex_to_RGB(self.panel_color_text) # 绿色
        lightGreenboxIte.setFlag(QGraphicsItem.ItemIsSelectable)  # 设置图形项可以选择
        self._scene.addItem(lightGreenboxIte)
        self.general_panel_showlist.append(lightGreenboxIte)
        self.showPanelList()

    def Hex_to_RGB(self,color):
        r = int(color[1:3], 16)
        g = int(color[3:5], 16)
        b = int(color[5:7], 16)

        rgb = str(r) + ',' + str(g) + ',' + str(b)
        return QColor(r, g, b)
    def getlSignal_refresh(self,lists):
        print("用于接收刷新信号")
        #删除面板
        if lists[0]==0:
            print("收到删除信号 进行局部刷新2  局部刷新")
            print("需要删除:", lists[1])
            self.showPanelList()
            for box in self.general_panel_showlist:
                if lists[1] == box.key:
                    self._scene.removeItem(box)
                    self.general_panel_showlist.remove(box)
                    print("删除完成",box.key)
        if lists[0]==2:
            print("修改编号之后刷新")
            #self.panel_info_all_list = sqlitefun.sqlite_select_panelinfolist(self.projecttime)
            self.showPanelList()
            if len(lists)>1:
                update_panel_info_list = lists[1]
                for box in self.general_panel_showlist:
                    for update_info in update_panel_info_list:
                        panel_id = update_info.split('|')[0]
                        panel_color = update_info.split('|')[1]
                        if box.key == panel_id:
                            box.color = self.Hex_to_RGB(panel_color)
            #self.showPanoramic()
        if lists[0]==3:
            print("完全刷新")
            #self.showPanelList()
            panoramic_bin = sqlitefun.sqlite_select_imagebin(self.projecttime, "panoramic")[0][0]
            self.showPanoramic(panoramic_bin)
        if lists[0]==4:
            print("批量删除刷新")
            self.showPanelList()
            for selectbox in lists[1]:
                for box in self.general_panel_showlist:
                    if selectbox == box.key:
                        self._scene.removeItem(box)
                        self.general_panel_showlist.remove(box)
                        print("删除完成",box.key)

    def showPanelList(self):
        self.listPanelID.clear()
        self.panel_info_dictlist.clear()
        self.panel_info_all_list = sqlitefun.sqlite_select_panelinfolist(self.projecttime)
        for panel_info_json in self.panel_info_all_list:
            panel_info_dict = json.loads(panel_info_json[1])  # 转为字典
            self.panel_info_dictlist.append(panel_info_dict)  # 加入字典列表方便使用
            # 列表显示增加
            self.listPanelID.addItem(panel_info_dict["panel_name"])

    def showMapImage(self):
        self._zoom = 0
        self._empty = True
        self._scene = RGISGraphicsScene() # 使用自定义的QGraphicsScene
        self._scene.signalMousePos.connect(self.selectPiont)
        #self._scene.selectItem.connect(self.selectPiont)
        self._photo = QtWidgets.QGraphicsPixmapItem()
        self._scene.addItem(self._photo)
        self.MapImage.setScene(self._scene)
        self.MapImage.setTransformationAnchor(QtWidgets.QGraphicsView.AnchorUnderMouse)
        self.MapImage.setResizeAnchor(QtWidgets.QGraphicsView.AnchorUnderMouse)
        self.MapImage.setVerticalScrollBarPolicy(QtCore.Qt.ScrollBarAlwaysOff)
        self.MapImage.setHorizontalScrollBarPolicy(QtCore.Qt.ScrollBarAlwaysOff)
        self.MapImage.setBackgroundBrush(QtGui.QBrush(QtGui.QColor(30, 30, 30)))
        self.MapImage.setFrameShape(QtWidgets.QFrame.NoFrame)
    
    def changeMoveMode(self):
        # 改变地图点击模式/移动模式
        if self.movemode == 0:
            # 现在是固定模式 变更到移动模式
            self.MapImage.setDragMode(QtWidgets.QGraphicsView.ScrollHandDrag)
            self.but_movemode.setText("拖动")
            self.movemode = 1
        else:
            self.MapImage.setDragMode(QtWidgets.QGraphicsView.NoDrag)
            self.but_movemode.setText("固定")
            self.movemode = 0

    
    def MapimageZoonOut(self):
        # 放大
         if self.hasPhoto():
            factor = 1.25
            self._zoom += 1

            if self._zoom == 0:
                self.fitInView()
            elif self._zoom <10:
                self.MapImage.scale(factor, factor)

            else:
                #self._zoom = 0
                self._zoom = 10
                self.te_bottom_hint.setText("已经放至最大了。")
            print("放大", self._zoom)
            
    def MapimageZoonIn(self):
        # 缩小
         if self.hasPhoto():
            factor = 0.75
            if self._zoom == 5:
                self._zoom=4
            self._zoom -= 1

            if self._zoom == 0:
                self.fitInView()
            elif self._zoom > int(-5):
                self.MapImage.scale(factor, factor)
            else:
                #self._zoom = 0
                self._zoom = -5
                self.te_bottom_hint.setText("已经放至最小了。")
            print("缩小", self._zoom)
    def MapimageZoonRest(self):
        # 还原
        #self.clickboxItem_list = []  # 刷新 清空选中列表
        self.getlSignal_refresh([3]) # 刷新 重载全部元素
    
    def hasPhoto(self):
        return not self._empty

    def fitInView(self, scale=True):
        rect = QtCore.QRectF(self._photo.pixmap().rect())
        if not rect.isNull():
            self.MapImage.setSceneRect(rect)
            if self.hasPhoto():
                unity = self.MapImage.transform().mapRect(QtCore.QRectF(0, 0, 1, 1))
                self.MapImage.scale(1 / unity.width(), 1 / unity.height())
                viewrect = self.MapImage.viewport().rect()
                scenerect = self.MapImage.transform().mapRect(rect)
                factor = min(viewrect.width() / scenerect.width(),
                             viewrect.height() / scenerect.height())
                self.MapImage.scale(factor, factor)
            self._zoom = 0
    def setPhoto(self, pixmap=None):
        self._zoom = 0
        if pixmap :
            self._empty = False
            
            self._photo.setPixmap(QtGui.QPixmap(QtGui.QImage.fromData(pixmap)))
        else:
            self._empty = True
            
            self._photo.setPixmap(QtGui.QPixmap())
        self.fitInView()
        
    def loadImage(self,mapimage_bin):
        self.setPhoto(mapimage_bin)

    def region_loc_alter(self):
        if self.bool_updage_number==1:
            self.te_bottom_hint.setText("正在修改组件编号，不可进行组件移动")
        elif self.bool_batch_update_panel == 1:
            self.te_bottom_hint.setText("正在批量修改组件信息，不可进行组件移动")
        else:
            #多面板位置移动
            if self.cb_label_mode.currentText() == "查看":
                print("多面板位置移动")
                self.btn_move_ok.setVisible(True)
                self.btn_move_no.setVisible(True)
                self.btn_batch_update_panel.setEnabled(False)
                self.btn_batch_del_panel.setEnabled(False)
                self.btn_GPS_range_computing.setEnabled(False)
                self.btn_update_number.setEnabled(False)
                self.btn_select_panel_info.setEnabled(False)
                self.btn_del_panel.setEnabled(False)
                print("是否允许面板移动",self.bool_region_move_panel)
                if self.bool_region_move_panel ==0:
                    self.bool_region_move_panel=1 #设置为1的时候  允许移动组件
                    self.MapImage.setDragMode(QGraphicsView.RubberBandDrag) #设置为可以拉框
                    self.btn_move_panel.setEnabled(False) #设置为不可点击
                    for box in self.general_panel_showlist:
                        box.bool_IsMovable = 1
                        box.setFlag(QGraphicsItem.ItemIsMovable,True)
            else:
                QMessageBox.information(self, "消息", "请将操作状态改为查看才可移动组件", QMessageBox.Yes)


    def start_fun_ok(self): #确定移动
        print("点击确定按钮")
        # 功能确认
        if self.bool_region_move_panel == 1:
            # 关闭区域选中 修改位置
            self.bool_region_move_panel = 0
            self.btn_move_panel.setEnabled(True)
            self.btn_batch_update_panel.setEnabled(True)
            self.btn_batch_del_panel.setEnabled(True)
            self.btn_GPS_range_computing.setEnabled(True)
            self.btn_update_number.setEnabled(True)
            self.btn_select_panel_info.setEnabled(True)
            self.btn_del_panel.setEnabled(True)
            self.btn_move_ok.setVisible(False)  # 隐藏掉功能按钮的取消键
            self.btn_move_no.setVisible(False)  # 隐藏掉功能按钮的确认键
            # 模式还原到使用之前
            if self.movemode == 1:
                self.MapImage.setDragMode(QtWidgets.QGraphicsView.ScrollHandDrag)
            else:
                self.MapImage.setDragMode(QtWidgets.QGraphicsView.NoDrag)

            box_moved_pos_dist = {}  # 保存移动完成后的坐标和组件编号
            sqlite_panel_info_all = []
            for box in self.general_panel_showlist:
                box.bool_IsMovable = 0
                # box.setFlag(QGraphicsItem.ItemIsMovable, False)  # 设置为不可移动状态
                box_x = str(int(box.sceneBoundingRect().topLeft().x()))
                box_y = str(int(box.sceneBoundingRect().topLeft().y()))
                box_moved_pos_dist.update({box.key: [box_x, box_y]})

                for panel_info in self.panel_info_all_list:
                    if panel_info[0] == box.key:
                        panel_info_dict = json.loads(panel_info[1])  # 转为字典
                        w = int(panel_info_dict["panel_wh"].split(',')[0])
                        h = int(panel_info_dict["panel_wh"].split(',')[1])
                        box_x = str(int(int(box_x)+(w/2)))
                        box_y = str(int(int(box_y)+(h/2)))
                        panel_info_dict["panel_pos"] = box_x+","+box_y


                        panel_info_json = json.dumps(panel_info_dict, ensure_ascii=False)
                        sqlite_panel_info_all.append([panel_info[0],panel_info_json])
            sqlitefun.sqlite_insert_panelinfo_all_json(self.projecttime,sqlite_panel_info_all)
            self.getlSignal_refresh([2])




    def end_fun_no(self): #取消移动
        print("点击取消")
        if self.bool_region_move_panel == 1:
            # 关闭区域选中 修改位置
            self.bool_region_move_panel = 0
            self.btn_move_panel.setEnabled(True)
            self.btn_batch_update_panel.setEnabled(True)
            self.btn_batch_del_panel.setEnabled(True)
            self.btn_GPS_range_computing.setEnabled(True)
            self.btn_update_number.setEnabled(True)
            self.btn_select_panel_info.setEnabled(True)
            self.btn_del_panel.setEnabled(True)
            self.btn_move_ok.setVisible(False)  # 隐藏掉功能按钮的取消键
            self.btn_move_no.setVisible(False)  # 隐藏掉功能按钮的确认键
            # 模式还原到使用之前
            if self.movemode == 1:
                self.MapImage.setDragMode(QtWidgets.QGraphicsView.ScrollHandDrag)
            else:
                self.MapImage.setDragMode(QtWidgets.QGraphicsView.NoDrag)
        box_moved_pos_dist = {}  # 保存移动完成后的坐标和组件编号
        for box in self.general_panel_showlist:
            box.bool_IsMovable = 0
            box.setFlag(QGraphicsItem.ItemIsMovable, False)  # 设置为不可移动状态
            box_x = str(int(box.intPos[0]))
            box_y = str(int(box.intPos[1]))
            box_moved_pos_dist.update({box.key: [box_x, box_y]})
            box.setPos(float(box_x), float(box_y))

    def selectPiont(self,arg1,arg2,mode):
        if (mode == 3):
            if len(arg1) > 0:
                self.selectItem = arg1



class RGISGraphicsScene(QGraphicsScene):
    # 自定义QGraphicsScene 场景 用于事件的传送
    selectItem = pyqtSignal(list)
    signalMousePos = pyqtSignal(list,list,int,int)
    #signalMousePos = pyqtSignal(list)

    def __init__(self,parent = None):
        super(RGISGraphicsScene, self).__init__(parent)
        self.pix =  QPixmap()  # 实例化一个 QPixmap 对象
        self.lastPoint =  QPoint() # 起始点
        self.endPoint =  QPoint() #终点    

    def mousePressEvent(self, event) :   
		# 鼠标左键按下
        super().mousePressEvent(event)
        if event.button() == Qt.LeftButton :
            self.lastPoint = event.scenePos()
            self.endPoint = self.lastPoint
            # print("libra按下坐标:",self.lastPoint,self.endPoint)
            #self.mapclickpos = [event.scenePos().x(), event.scenePos().y(), self.selectedItems()]
            #self.mapclickpos = self.selectedItems()
            self.signalMousePos.emit(self.selectedItems(),[],3,0)
        if event.button() == Qt.RightButton:
            self.signalMousePos.emit(self.selectedItems(), [],3, 1)

            
	# 鼠标移动事件
    def mouseMoveEvent(self, event):	
        # 鼠标左键按下的同时移动鼠标
        super().mouseMoveEvent(event)
        if event.buttons() and Qt.LeftButton :
            self.endPoint = event.scenePos()
            self.signalMousePos.emit([self.lastPoint.x(),self.lastPoint.y()],[self.endPoint.x(),self.endPoint.y()],1,0)
            # print("libra移动坐标:",self.lastPoint,self.endPoint)
            #进行重新绘制


    # 鼠标释放事件
    def mouseReleaseEvent( self, event):
		# 鼠标左键释放
        super().mouseReleaseEvent(event)
        if event.button() == Qt.LeftButton :
            self.endPoint = event.scenePos()
            self.signalMousePos.emit([self.lastPoint.x(),self.lastPoint.y()],[self.endPoint.x(),self.endPoint.y()],0,0)
            # print("libra释放坐标:",self.lastPoint,self.endPoint)
			#进行重新绘制
    #def

class PanelrectboxItem(QGraphicsItem):
 # 面板矩形框填充
    def __init__(self,key):
        super(PanelrectboxItem, self).__init__()
        self.color = Qt.blue     # 显示面板的颜色
        self.rectpos = [0,0,100,100]  # 显示面板的坐标
        self.key = key
        self.name = "-1,-1"
        self.bool_IsMovable = 0
        self.intPos = [self.rectpos[0], self.rectpos[1], self.rectpos[2], self.rectpos[3]]



    def boundingRect(self):
        # return PanelrectboxItem.Rect
        return QRectF(self.rectpos[0],self.rectpos[1],self.rectpos[2],self.rectpos[3])
    def resname(self):
        # 用于返回key名称
        return self.key
 
    def paint(self, painter, option, widget):
        # 画个矩形框
        if self.isSelected():
            painter.setPen(Qt.white)  # 选中的颜色
            painter.setBrush(Qt.yellow)
        else:
            painter.setPen(self.color)
            painter.setBrush(self.color)
        painter.drawRect(self.rectpos[0],self.rectpos[1],self.rectpos[2],self.rectpos[3])
        if self.bool_IsMovable == 1:
            self.setFlag(QGraphicsItem.ItemIsMovable, True)  #  故障组件设置为可移动状态
        if self.bool_IsMovable == 0:
            self.setFlag(QGraphicsItem.ItemIsMovable, False)

    def mousePressEvent(self,event):
        # self.mapclickpos = [event.pos().x(),event.pos().y()]
        # print(self.mapclickpos)
        super().mousePressEvent(event)  # 将点击坐标传递下去


class ImagerectboxItem_wh(QGraphicsItem):
 # 面板矩形框
    def __init__(self,PlanetType):
        super(ImagerectboxItem_wh, self).__init__()
        self.color = Qt.blue     # 显示面板的颜色
        self.rectpos = [0,0,0,0]  # 显示面板的坐标
        self.type = PlanetType
        self.boundingRect = [self.rectpos[0],self.rectpos[1],self.rectpos[2],self.rectpos[3]]
        
    def boundingRect(self):
        return QRectF(self.rectpos[0],self.rectpos[1],self.rectpos[2],self.rectpos[3])
 
    def paint(self, painter, option, widget):
        # 画个矩形框
        painter.setPen(self.color)
        painter.drawRect(self.rectpos[0],self.rectpos[1],self.rectpos[2],self.rectpos[3])

class OpenProject(QMainWindow,Ui_open_project):
    # 打开选择历史UI界面
    OpenProject_Signal = pyqtSignal(list) # 自定义信号
    def __init__(self,parent = None):
        super(OpenProject, self).__init__(parent)
        self.setupUi(self)
        self.setFixedSize(self.width(), self.height())
        self.projectlist = sqlitefun.sqlite_select_localist()
        for prlist in self.projectlist:
            self.lw_local_PN.addItem(prlist[0]+ ' / '+prlist[1])
        self.lw_local_PN.setContextMenuPolicy(Qt.CustomContextMenu)
        self.lw_local_PN.customContextMenuRequested.connect(self.select_local_menu)  # 绑定列表
        self.but_open.clicked.connect(self.butOpenClicked)

    def setProjectInfo(self,projectTime):
        self.old_projectTime = projectTime

    def select_local_menu(self):
        select_item = self.lw_local_PN.currentItem()
        project = select_item.text()
        project_time = project.split(' / ')[0]
        menu = QMenu()
        opt1 = menu.addAction("删除")
        action = menu.exec_(QtGui.QCursor.pos())
        if action == opt1:
            if self.old_projectTime == project_time:
                QMessageBox.question(self, "提示", project + "已被打开，不能删除此项目", QMessageBox.Yes)

            else:
                reply = QMessageBox.question(self, "提示", "你确定要项目:" + project + "?",
                                             QMessageBox.Yes | QMessageBox.No)
                if reply == QMessageBox.Yes:

                    self.lw_local_PN.takeItem(self.lw_local_PN.row(select_item))
                    old_template_list = []
                    old_panel_info_list = sqlitefun.sqlite_select_panelinfolist(project_time)
                    for old_panel_info in old_panel_info_list:
                        old_panel_dict = json.loads(old_panel_info[1])
                        old_temple = old_panel_dict["panel_model"]
                        old_template_list.append(old_temple)
                    old_template_list = set(old_template_list)
                    for delete_model in old_template_list:
                        update_use_projectid = ""
                        result = sqlitefun.sqlite_select_use_template(delete_model)
                        projectid_list = result[0][0].split(';')
                        for projectid in projectid_list:
                            if projectid != project_time and projectid != '':
                                update_use_projectid += projectid + ";"
                        sqlitefun.sqlite_update_use_template(delete_model, update_use_projectid)
                    sqlitefun.sqlite_delete_panelinfo(project_time)
                    print(project_time)
    def butOpenClicked(self):
        # 历史列表打开
        checkItems = self.lw_local_PN.selectedItems()
        if checkItems:
            self.project_time = str(checkItems[0].text()).split(" / ")[0]
            self.project_name = str(checkItems[0].text()).split(" / ")[1]
            self.send_main_Signal()
            self.close()

    def send_main_Signal(self):
        print("发送相关信息给主界面")
        lists = ["OpenProject",self.project_time,self.project_name,self.old_projectTime]
        self.OpenProject_Signal.emit(lists) # 发射信号 


class NumDialog(QMainWindow,Ui_numDialog):
    NumDialog_Signal = pyqtSignal(list)  # 自定义信号
    def __init__(self,parent = None):
        super(NumDialog, self).__init__(parent)
        self.setupUi(self)
        self.setFixedSize(self.width(), self.height())
        self.Ok_button.clicked.connect(self.butOkClicked)
        self.Cancel_button.clicked.connect(self.butcancelClicked)


    def butOkClicked(self):
        self.numValue = self.Num_value.text()
        if self.numValue == '':
            QMessageBox.information(self, "消息", "行数量不能为空，请重新输入", QMessageBox.Yes)
        else:
            self.send_main_Signal()
            self.close()
    def butcancelClicked(self):
        self.numValue = -1
        self.send_main_Signal()
        self.close()
    def send_main_Signal(self):
        print("发送相关信息给主界面")
        lists = ["NumDialog", self.numValue]
        self.NumDialog_Signal.emit(lists)  # 发射信号

class Row_Column_Dialog(QMainWindow, Ui_row_Column_Dialog):
    Row_Column_Dialog_Signal = pyqtSignal(list)  # 自定义信号

    def __init__(self, parent=None):
        super(Row_Column_Dialog, self).__init__(parent)
        self.setupUi(self)
        self.setFixedSize(self.width(), self.height())
        self.ok_row_column_btn.clicked.connect(self.butOkrcClicked)
        self.cancel_row_column_btn.clicked.connect(self.butCaneclrcClicked)
    def butCaneclrcClicked(self):
        self.row_value = -1
        self.column_value = -1
        self.send_main_Signal()
        self.close()
    def butOkrcClicked(self):
        self.row_value = self.row_num_value.text()
        self.column_value = self.column_num_value.text()
        if self.row_value=='':
            QMessageBox.information(self, "消息", "行数量不能为空，请重新输入", QMessageBox.Yes)
        elif self.column_value=='':
            QMessageBox.information(self, "消息", "列数量不能为空，请重新输入", QMessageBox.Yes)
        else:
            self.send_main_Signal()
            self.close()

    def send_main_Signal(self):
        print("发送相关信息给主界面")
        lists = ["Row_Column_Dialog", self.row_value,self.column_value]
        self.Row_Column_Dialog_Signal.emit(lists)  # 发射信号

        #self.OK_button.clicked.connect(self.OK_buttonClicked)
class SetPanelInfo(QMainWindow,Ui_setPanelInfo):
    # 设置组件信息
    SetPanelInfo_Signal = pyqtSignal(list) # 自定义信号
    def __init__(self,parent = None):
        super(SetPanelInfo, self).__init__(parent)
        self.setupUi(self)
        self.setFixedSize(self.width(), self.height())
        sqlitefun.sqlite_create_template()
        self.new_label_type.addItem("光伏组件")
        self.panel_wh_text = sqlitefun.sqlite_select_config("panel_wh")
        if len(self.panel_wh_text) == 0:
            self.panel_wh_text = "10,10"
        else:
            self.panel_wh_text = self.panel_wh_text[0][0]
        self.new_panel_wh.setText(self.panel_wh_text)

        self.color=None
        self.but_new_save.clicked.connect(self.butNewSaveClicked)  # 保存创建信息
        #self.but_old_save.clicked.connect(self.butOldSaveClicked) #选择模板信息
        self.btn_panel_color.clicked.connect(self.butPanelColorClicked)#x选择颜色信息


    def butNewSaveClicked(self):

        panel_Template = self.new_panel_model.text()
        self.panel_Template_text = sqlitefun.sqlite_select_template(panel_Template)
        if len(self.panel_Template_text)>0:
            QMessageBox.information(self, "消息", "模板名称已存在，请重新输入模板名称", QMessageBox.Yes)
            print("模板名称已存在，请重新输入模板名称")
        elif self.color == None:
            QMessageBox.information(self, "消息", "请选择组件颜色", QMessageBox.Yes)
            print("请选择组件颜色")
        else:
            panel_type = self.new_label_type.currentText() #标注类型
            panel_manufactor = self.new_label_manufactor.text()#组件厂家
            panel_model = self.new_panel_model.text()#组件型号
            panel_power = self.new_panel_power.text()#组件功率.
            panel_wh = self.new_panel_wh.text()#组件宽高
            panel_altitude = self.new_panel_altitude.text()#组件海拔
            panel_angle = self.new_panel_angle.text()#组件角度
            panel_arrange = self.new_panel_arrange.text()#组件排布
            panel_color = self.color
            if panel_arrange == "":
                QMessageBox.information(self, "消息", "组件排布不能为空，格式为2x10", QMessageBox.Yes)
            else:
                IsFormat = panel_arrange.split("x")
                boo = False
                if (len(IsFormat) == 2):
                    if self.is_number(IsFormat[0]) and self.is_number(IsFormat[1]):
                        boo = True
                if boo:
                    dict_value = {'panel_manufactor': panel_manufactor,
                                  'label_type': panel_type, 'panel_model': panel_model,
                                  'panel_power': panel_power, 'panel_wh': panel_wh,
                                  'panel_altitude': panel_altitude,
                                  'panel_angle': panel_angle, 'panel_arrange': panel_arrange,
                                  'panel_color': panel_color
                                  }
                    panel_json = json.dumps(dict_value, ensure_ascii=False)
                    sqlitefun.sqlite_insert_template(panel_Template, panel_json)
                    print("模板名称", panel_Template)
                    # print("wqp",self.dict_value)
                    # 保存上次记录数据
                    sqlitefun.sqlite_insert_config("label_type", panel_type)
                    sqlitefun.sqlite_insert_config("panel_model", panel_model)
                    sqlitefun.sqlite_insert_config("panel_power", panel_power)
                    sqlitefun.sqlite_insert_config("panel_wh", panel_wh)
                    sqlitefun.sqlite_insert_config("panel_altitude", panel_altitude)
                    sqlitefun.sqlite_insert_config("panel_angle", panel_angle)
                    sqlitefun.sqlite_insert_config("panel_manufactor", panel_manufactor)
                    sqlitefun.sqlite_insert_config("panel_arrange", panel_arrange)
                    sqlitefun.sqlite_insert_config("panel_color", panel_color)

                    self.send_main_Signal1()
                    self.close()
                else:
                    QMessageBox.question(self, "错误提示", "输入的组件排布格式错误，请重新输入组件排布，格式为 2x10", QMessageBox.Yes)



    #选择模板后确定信息
    # def butOldSaveClicked(self):
    #     template_name = self.old_panel_model.currentText()
    #     print("模型名称", template_name)
    #     panel_json = sqlitefun.sqlite_select_template(template_name)
    #     panel_dict = json.loads(panel_json[0][0])
    #     sqlitefun.sqlite_insert_config("label_type", panel_dict["label_type"])
    #     sqlitefun.sqlite_insert_config("panel_model", panel_dict["panel_model"])
    #     sqlitefun.sqlite_insert_config("panel_power", panel_dict["panel_power"])
    #     sqlitefun.sqlite_insert_config("panel_wh", panel_dict["panel_wh"])
    #     sqlitefun.sqlite_insert_config("panel_altitude", panel_dict["panel_altitude"])
    #     sqlitefun.sqlite_insert_config("panel_angle", panel_dict["panel_angle"])
    #     sqlitefun.sqlite_insert_config("panel_manufactor", panel_dict["panel_manufactor"])
    #     sqlitefun.sqlite_insert_config("panel_arrange", panel_dict["panel_arrange"])
    #     sqlitefun.sqlite_insert_config("panel_color", panel_dict["panel_color"])
    #     self.close()
    def butPanelColorClicked(self):
        print("选择颜色信息")
        col = QColorDialog.getColor()
        self.new_panel_color.setStyleSheet('QWidget {background-color:%s}' % col.name())
        self.color = col.name()
    def is_number(self,s):
        try:
            int(s)
            return True
        except ValueError:
            pass
        return False
    def send_main_Signal1(self):
        print("发送相关信息给主界面")
        lists = ["SetPanel_Info","1111"]
        self.SetPanelInfo_Signal.emit(lists) # 发射信号

class CreateProject(QMainWindow,Ui_creatProject):
    createproject_Signal = pyqtSignal(list) # 自定义信号
    def __init__(self,parent = None):
        super(CreateProject, self).__init__(parent)
        self.setupUi(self)
        self.setFixedSize(self.width(), self.height())
        self.project_time = datetime.datetime.now().strftime('%Y%m%d_%H%M%S_%f')   #  获取毫秒时间戳
        self.projectName = ""
        self.lab_projectTime.setText("项目创建时间:"+self.project_time)

        self.but_create.clicked.connect(self.butCreateClicked)
        self.but_map_path.clicked.connect(self.openMapPath)
        sqlitefun.sqlite_init_execute()  # 初始化项目信息表

    def setProjectInfo(self,projectTime):
        self.old_projectTime = projectTime
        
    
    def butCreateClicked(self):
        # 相应创建项目按钮事件
        self.projectName = self.le_projectName.text()
        self.mappath = self.le_map_path.text()
        
        if self.projectName == "":
            return False
        sqlitefun.sqlite_create_projecttime(self.project_time,self.projectName)  # 向数据库写入项目时间和项目名称
        basicsfun.savaImagebinSqlite(self.project_time,"panoramic",self.mappath) # 全景图片写入数据库
        print("创建项目名称:"+self.projectName)
        self.send_main_Signal()
        self.close()
    
    def send_main_Signal(self):
        print("发送相关信息给主界面")
        lists = ["CreateProject",self.project_time,self.projectName,self.old_projectTime]
        self.createproject_Signal.emit(lists) # 发射信号
        
    
    def openMapPath(self):
        filename,  _ = QFileDialog.getOpenFileName(self, '请选择地图文件', None,"*.jpg")
        self.le_map_path.setText(filename)

class Update_Panel_info(QMainWindow,Ui_Upadte_Panel_info):
    Update_Panel_info_Signal = pyqtSignal(list)  # 自定义信号
    def __init__(self, parent=None):
        super(Update_Panel_info, self).__init__(parent)
        self.setupUi(self)
        self.setFixedSize(self.width(), self.height())
        self.Ok_button.clicked.connect(self.butOkClicked)

    def butOkClicked(self):
        self.number_value = self.new_number_value.text()
        self.send_main_Signal()

    def setT(self,number_value,item):
        self.new_number_value.setText(number_value)
        self.listItem = item
    def send_main_Signal(self):
        print("发送相关信息给主界面")
        lists = ["Update_Panel_info", self.number_value,self.listItem]
        self.Update_Panel_info_Signal.emit(lists)  # 发射信号

class Update_Number_dialog(QMainWindow,Ui_Update_number_Dialog):
    Update_Number_dialog_Signal = pyqtSignal(list)  # 自定义信号
    def __init__(self, parent=None):
        super(Update_Number_dialog, self).__init__(parent)
        self.setupUi(self)
        self.setFixedSize(self.width(), self.height())

        self.Ok_button.clicked.connect(self.butOkClicked)
        self.Cancel_button.clicked.connect(self.butCanelClicked)

    def butOkClicked(self):
        self.number_value = self.Number_value.text()
        self.direction_value = self.number_direction.currentText()
        boo = False
        IsFormat = self.number_value.split(',')
        if (len(IsFormat) == 3):
            if self.is_number(IsFormat[0]) and self.is_number(IsFormat[1]):
                boo = True
        if boo:
            self.send_main_Signal()
        else:
            QMessageBox.question(self, "错误提示", "输入的格式错误，请重新输入编号，格式为 1,1,x", QMessageBox.Yes)

    def butCanelClicked(self):
        self.number_value=0
        self.send_main_Signal()
        self.close()
    def is_number(self,s):
        try:
            int(s)
            return True
        except ValueError:
            pass
        return False
    # def setT(self,number_value,item):
    #     self.new_number_value.setText(number_value)
    #     self.listItem = item
    def send_main_Signal(self):
        print("发送相关信息给主界面")
        lists = ["Update_Number_dialog", self.number_value,self.direction_value]
        self.Update_Number_dialog_Signal.emit(lists)  # 发射信号

class Select_Update_Panel_info(QMainWindow,Ui_update_Panel):
    Select_Update_Panel_info_Signal = pyqtSignal(list)  # 自定义信号
    def __init__(self, parent=None):
        super(Select_Update_Panel_info, self).__init__(parent)
        self.setupUi(self)
        self.setFixedSize(self.width(), self.height())
        self.listItem = None
        self.batch_list = []
        self.dict_value = None
        self.IsUpdate = -1
        #self.panel_info_color = None

        self.but_save.clicked.connect(self.butOkClicked)
        self.btu_close.clicked.connect(self.butNoClicked)
    def butNoClicked(self):
        self.IsUpdate = 0
        self.send_main_Signal()
    def butOkClicked(self):
        panel_name_value = self.panel_info_name.text()
        panel_manufactor_value = self.panel_info_manufactor.text()
        panel_type_value = self.panel_info_type.text()
        panel_model_value = self.panel_info_model.currentText()
        panel_power_value = self.panel_info_power.text()
        panel_wh_value = self.panel_info_wh.text()
        panel_altitude_value = self.panel_info_altitude.text()
        panel_angle_value = self.panel_info_angle.text()
        panel_arrange_value = self.panel_info_arrange.text()
        panel_gps_value = self.panel_info_GPS.text()
        panel_color = self.panel_info_color
        gps_value = panel_gps_value.split(',')
        longitude = gps_value[0]  # 经度
        latitude = gps_value[1]  # 纬度
        listItem = self.listItem
        if listItem != None:
            itemvalue = listItem.text()
            isItem = False;
            listcount = self.listPanelID.count()

            if panel_name_value != itemvalue:
                for i in range(listcount):
                    if panel_name_value == self.listPanelID.item(i).text():
                        isItem = True;
                        break;
                    else:
                        isItem = False;
            if float(longitude) > 180 or float(longitude) < -180:
                QMessageBox.information(self, "消息", "经度值错误(不得大于180或者不得小于-180)，请重新输入", QMessageBox.Yes)
            elif float(latitude) > 90 or float(latitude) < -90:
                QMessageBox.information(self, "消息", "纬度值错误(不得大于90或者不得小于-90)，请重新输入", QMessageBox.Yes)
            elif isItem :
                QMessageBox.information(self, "消息", "编号已存在，请重新输入", QMessageBox.Yes)
            else:
                print("修改颜色",panel_color)
                self.dict_value = {'panel_name':panel_name_value,'panel_manufactor':panel_manufactor_value,'label_type':panel_type_value,'panel_model':panel_model_value,
                                   'panel_power':panel_power_value,'panel_wh':panel_wh_value,'panel_altitude':panel_altitude_value,
                                   'panel_angle':panel_angle_value,'panel_arrange':panel_arrange_value,'panel_gps':panel_gps_value,
                                   'panel_color':panel_color
                                   }
        else:
            self.dict_value = {'panel_name': panel_name_value, 'panel_manufactor': panel_manufactor_value,
                               'label_type': panel_type_value, 'panel_model': panel_model_value,
                               'panel_power': panel_power_value, 'panel_wh': panel_wh_value,
                               'panel_altitude': panel_altitude_value,
                               'panel_angle': panel_angle_value, 'panel_arrange': panel_arrange_value,
                               'panel_gps': panel_gps_value,
                               'panel_color': panel_color
                               }
        self.send_main_Signal()

    def setT(self,panel_dict,item,listPanelID):
        # self.panel_info_type.addItem("光伏组件2")
        # self.panel_info_type.addItem("光伏组件")
        self.listPanelID = listPanelID
        model_index = -1
        self.panel_info_type.setText(panel_dict["label_type"])

        self.panel_info_name.setText(panel_dict["panel_name"])
        self.panel_info_manufactor.setText(panel_dict["panel_manufactor"])

        template_list = sqlitefun.sqlite_select_template_all()
        for templat in template_list:
            self.panel_info_model.addItem(templat[0])
        for count in range(self.panel_info_model.count()):
            if self.panel_info_model.itemText(count) == panel_dict["panel_model"]:
                model_index = count
                break


        self.panel_info_model.setCurrentIndex(model_index)
        self.panel_info_model.currentIndexChanged.connect(self.modelBoxClick)
        self.panel_info_power.setText(panel_dict["panel_power"])
        self.panel_info_wh.setText(panel_dict["panel_wh"])
        self.panel_info_altitude.setText(panel_dict["panel_altitude"])
        self.panel_info_angle.setText(panel_dict["panel_angle"])
        self.panel_info_arrange.setText(panel_dict["panel_arrange"])
        self.panel_info_GPS.setText(panel_dict["panel_gps"])
        self.panel_info_color = panel_dict["panel_color"]
        self.listItem = item
    def modelBoxClick(self):
        print("组件模型：",self.panel_info_model.currentText())
        panel_model = self.panel_info_model.currentText()
        self.panel_json = sqlitefun.sqlite_select_template(panel_model)
        panel_dict = json.loads(self.panel_json[0][0])

        self.panel_info_manufactor.setText(panel_dict["panel_manufactor"])
        self.panel_info_type.setText(panel_dict["label_type"])
        # self.old_panel_model.setText(panel_dict["panel_model"])
        self.panel_info_power.setText(panel_dict["panel_power"])
        self.panel_info_wh.setText(panel_dict["panel_wh"])
        self.panel_info_altitude.setText(panel_dict["panel_altitude"])
        self.panel_info_angle.setText(panel_dict["panel_angle"])
        self.panel_info_arrange.setText(panel_dict["panel_arrange"])
        self.panel_info_color = panel_dict["panel_color"]

    def set_batch_update(self,batch_update_list):
        self.panel_info_name.setEnabled(False)
        self.panel_info_manufactor.setEnabled(False)
        self.panel_info_type.setEnabled(False)
        self.panel_info_model.setEnabled(True)
        self.panel_info_model.currentIndexChanged.connect(self.modelBoxClick)
        self.panel_info_power.setEnabled(False)
        self.panel_info_wh.setEnabled(False)
        self.panel_info_altitude.setEnabled(False)
        self.panel_info_angle.setEnabled(False)
        self.panel_info_arrange.setEnabled(False)
        self.panel_info_GPS.setEnabled(False)
        self.batch_list = batch_update_list
    def send_main_Signal(self):
        lists = ["Select_Update_Panel_info", self.dict_value,self.listItem,self.batch_list,self.IsUpdate]
        self.Select_Update_Panel_info_Signal.emit(lists)  # 发射信号
class GD_map_Window(QMainWindow):
    def __init__(self):
        super(GD_map_Window, self).__init__()
        self.setWindowTitle('google地图显示')  # 窗口标题
        self.resize(981,618)

    def setUrl(self,url):
        self.browser = QWebEngineView()
        self.browser.load(QUrl(url))
        self.setCentralWidget(self.browser)
class Upload_Chain_class(QMainWindow,Ui_Chain_object):
    Upload_Chain_Signal = pyqtSignal(list)  # 自定义信号

    def __init__(self, parent=None):
        super(Upload_Chain_class, self).__init__(parent)
        self.setupUi(self)
        self.setFixedSize(self.width(), self.height())
        self.btn_upload.clicked.connect(self.upload_chain)#上传按钮
        self.btu_close.clicked.connect(self.upload_close)#取消按钮

    def setT(self,local_project_list,chain_project_list):
        self.local_project_info_list = local_project_list
        for local_project in local_project_list:
            project_time = local_project[0]
            project_name = local_project[1]
            self.local_list.addItem(project_time+"/"+project_name)

        for chain_project in chain_project_list:
            self.chain_list.addItem(chain_project)


    def upload_close(self):
        print("取消按钮")
        self.close()
    def upload_chain(self):
        #开启进度条
        loading = Loading_class()
        loading.show()
        #弹出框文字提示，并且不可关闭
        checkItems = self.local_list.selectedItems()

        class DownloadTaskHandler(threading.Thread):
            def run(self):
                print("上传到云端数据库")

                select_project = str(checkItems[0].text())
                select_project_time = str(checkItems[0].text()).split("/")[0]
                project_name = str(checkItems[0].text()).split("/")[1]

                # 初始化区块链信息
                token = RequestsPost.getToken()
                plant_init_result = RequestsPost.plant_chain_init(select_project_time, token)
                # plant_init_result = RequestsPost.plant_chain_init("888", token)
                print("plant_init_result", plant_init_result)
                plant_root_table_name = ""
                chain_hand_id = ""  # 头id
                chain_tail_id = ""  # 尾id
                chain_type = "plant_root"  # 类型
                chain_num = 0  # 数量
                creat_id = "test_id"
                json_data = ""
                if plant_init_result != 0:
                    for plant_init_info in plant_init_result:
                        plannt_type = plant_init_info["chain_type"]
                        if plannt_type == "plant_root":
                            plant_root_table_name = plant_init_info["tabel_name"]
                            # chain_hand_id = plant_init_info["chain_hand"]
                    plant_table_info_result = RequestsPost.plant_chain_info_select(plant_root_table_name, token)

                    # chain_tail_id = plant_init_info["chain_father"]

                else:
                    # 如果已存在，查询plant_table
                    plant_table_info_result = RequestsPost.plant_chain_info_select("plant_table", token)
                    for plant_table_info in plant_table_info_result:
                        print("plant_table_info", plant_table_info)
                        plant_type = plant_table_info["chain_type"]
                        plant_id = plant_table_info["plant_id"]
                        if plant_id == select_project_time:
                            if plant_type == "plant_root":
                                plant_root_table_name = plant_table_info["tabel_name"]

                print("plant_table_name", plant_root_table_name)
                plant_root_info = RequestsPost.plant_chain_info_select(plant_root_table_name, token)

                chain_num = len(plant_root_info)
                if chain_num == 0:
                    chain_tail_id = "root"
                else:
                    chain_hand_id = plant_root_info[-1]["chain_root_hand"]
                    chain_tail_id = plant_root_info[-1]["chain_id"]
                print("链表的链数和", len(plant_root_info))

                # 获取所有组件信息
                panel_info_all_list = sqlitefun.sqlite_select_panelinfolist(select_project_time)
                panel_info_all_json_list = []
                panel_model_json_list = []
                # 获取当前项目使用的模板信息
                panel_model_all = []
                panel_model_list = []
                for panel_info_all in panel_info_all_list:
                    panel_info_all_dict = json.loads(panel_info_all[1])
                    panel_model_all.append(panel_info_all_dict['panel_model'])
                    panel_info_all_json_list.append(panel_info_all[1])

                [panel_model_list.append(i) for i in panel_model_all if not i in panel_model_list]

                for panel_model in panel_model_list:
                    # 查询模板详细信息
                    panel_model_json = sqlitefun.sqlite_select_template(panel_model)
                    panel_model_dict = json.loads(panel_model_json[0][0])
                    panel_model_json_list.append(panel_model_dict)

                print("panel_model_json", panel_model_json_list)
                # 获取图片
                panoramic_bin = sqlitefun.sqlite_select_imagebin(select_project_time, "panoramic")[0][0]
                image_byte_list = []
                image_name_list = []
                is_panoramic = True
                num = 0
                lens = 0
                print(len(panoramic_bin))
                while is_panoramic:
                    lens = num
                    num = lens + 512000
                    if num > (len(panoramic_bin) - 1):
                        num = len(panoramic_bin) - 1
                        is_panoramic = False

                    bst = panoramic_bin[lens:num]
                    icon_byte = base64.b64encode(bst)
                    icon_str = icon_byte.decode('ascii')
                    image_name = RequestsPost.image_upload(icon_str, token)
                    print(image_name)
                    image_name_list.append(image_name)
                print(image_name_list)

                #  byte类型转换为str
                json_dic = {'project_id': select_project_time, 'project_name': project_name,
                            'panel_info_all_json': panel_info_all_json_list, 'panel_model_json': panel_model_json_list,
                            'panel_image_list': image_name_list}
                json_data = json.dumps(json_dic, ensure_ascii=False)
                # print("json_data",json_data)
                RequestsPost.plant_chain_add_info(token, select_project_time, chain_hand_id, chain_tail_id, chain_type,
                                                  chain_num, creat_id, json_data)
                loading.close()


        DownloadTaskHandler(daemon=True).start()
        #隐藏进度条

    def send_main_Signal(self):
        lists = ["Upload_Chain_Signal"]
        self.Upload_Chain_Signal.emit(lists)  # 发射信号

class Loading_class(QMainWindow,Ui_Loading_object):
    def __init__(self, parent=None):
        super(Loading_class, self).__init__(parent)
        self.setupUi(self)
        self.setFixedSize(self.width(), self.height())


class Update_Panel_Model_class(QMainWindow,Ui_Update_Panel_Model):
    Update_Panel_Model_Signal = pyqtSignal(list)  # 自定义信号

    def __init__(self, parent=None):
        super(Update_Panel_Model_class, self).__init__(parent)
        self.setupUi(self)
        self.setFixedSize(self.width(), self.height())
        self.panel_m_type.addItem("光伏组件")
        self.btn_panel_m_color.clicked.connect(self.btn_Panel_Color_Clicked)
        self.but_new_save.clicked.connect(self.btn_update_model_Clicked)
        self.but_delete.clicked.connect(self.btn_delede_model_Clicked)
        self.color=""
    def setT(self,type):
        if type==1:
            #修改
            self.but_new_save.setVisible(True)
            self.but_delete.setVisible(False)
        elif type ==2:
            #删除
            self.but_new_save.setVisible(False)
            self.but_delete.setVisible(True)
            self.panel_m_type.setEnabled(False)
            # self.old_panel_model.setText(panel_dict["panel_model"])
            self.panel_m_manufactor.setEnabled(False)
            self.panel_m_power.setEnabled(False)
            self.panel_m_wh.setEnabled(False)
            self.panel_m_altitude.setEnabled(False)
            self.panel_m_angle.setEnabled(False)
            self.panel_m_arrange.setEnabled(False)
            self.btn_panel_m_color.setEnabled(False)

        template_list = sqlitefun.sqlite_select_template_all()
        old_model = sqlitefun.sqlite_select_config("panel_model")
        if len(template_list) > 0:
            item_index = 0
            for templat in template_list:
                self.panel_model.addItem(templat[0])
            for count in range(self.panel_model.count()):
                if old_model[0][0] == self.panel_model.itemText(count):
                    item_index = count
                    break
            self.panel_model.setCurrentIndex(item_index)

            panel_json = sqlitefun.sqlite_select_template(old_model[0][0])
            panel_dict = json.loads(panel_json[0][0])

            self.panel_m_manufactor.setText(panel_dict["panel_manufactor"])
            for count in range(self.panel_m_type.count()):
                if panel_dict["label_type"] == self.panel_m_type.itemText(count):
                    item_index = count
                    break
            self.panel_m_type.setCurrentIndex(item_index)
            #self.old_panel_model.setText(panel_dict["panel_model"])
            self.panel_m_power.setText(panel_dict["panel_power"])
            self.panel_m_wh.setText(panel_dict["panel_wh"])
            self.panel_m_altitude.setText(panel_dict["panel_altitude"])
            self.panel_m_angle.setText(panel_dict["panel_angle"])
            self.panel_m_arrange.setText(panel_dict["panel_arrange"])
            self.panel_m_color.setStyleSheet('QWidget {background-color:%s}' % panel_dict["panel_color"])
            self.color = panel_dict["panel_color"]
            self.panel_model.currentIndexChanged.connect(self.panel_model_Clicked)

    def panel_model_Clicked(self):
        print("选择组件型号进行切换")
        panel_model_name = self.panel_model.currentText()
        panel_json = sqlitefun.sqlite_select_template(panel_model_name)
        panel_dict = json.loads(panel_json[0][0])
        self.panel_m_manufactor.setText(panel_dict["panel_manufactor"])
        for count in range(self.panel_m_type.count()):
            if panel_dict["label_type"] == self.panel_m_type.itemText(count):
                item_index = count
                break
        self.panel_m_type.setCurrentIndex(item_index)
        # self.old_panel_model.setText(panel_dict["panel_model"])
        self.panel_m_power.setText(panel_dict["panel_power"])
        self.panel_m_wh.setText(panel_dict["panel_wh"])
        self.panel_m_altitude.setText(panel_dict["panel_altitude"])
        self.panel_m_angle.setText(panel_dict["panel_angle"])
        self.panel_m_arrange.setText(panel_dict["panel_arrange"])
        self.panel_m_color.setStyleSheet('QWidget {background-color:%s}' % panel_dict["panel_color"])
        self.color = panel_dict["panel_color"]
    def is_number(self,s):
        try:
            int(s)
            return True
        except ValueError:
            pass
        return False
    def btn_Panel_Color_Clicked(self):
        print("修改模板颜色")
        col = QColorDialog.getColor()
        self.panel_m_color.setStyleSheet('QWidget {background-color:%s}' % col.name())
        self.color = col.name()
    def btn_update_model_Clicked(self):
        print("修改模板信息")
        #将修改后的模板信息保存到模板信息表中
        panel_model = self.panel_model.currentText()#组件型号
        panel_type = self.panel_m_type.currentText()  # 标注类型
        panel_manufactor = self.panel_m_manufactor.text()  # 组件厂家
        panel_power = self.panel_m_power.text()  # 组件功率.
        panel_wh = self.panel_m_wh.text()  # 组件宽高
        panel_altitude = self.panel_m_altitude.text()  # 组件海拔
        panel_angle = self.panel_m_angle.text()  # 组件角度
        panel_arrange = self.panel_m_arrange.text()  # 组件排布
        panel_color = self.color
        if panel_arrange == "":
            QMessageBox.information(self, "消息", "组件排布不能为空，格式为2x10", QMessageBox.Yes)
        else:
            IsFormat = panel_arrange.split("x")
            boo = False
            if (len(IsFormat) == 2):
                if self.is_number(IsFormat[0]) and self.is_number(IsFormat[1]):
                    boo = True
            if boo:
                dict_value = {'panel_manufactor': panel_manufactor,
                              'label_type': panel_type, 'panel_model': panel_model,
                              'panel_power': panel_power, 'panel_wh': panel_wh,
                              'panel_altitude': panel_altitude,
                              'panel_angle': panel_angle, 'panel_arrange': panel_arrange,
                              'panel_color': panel_color
                              }
                panel_json = json.dumps(dict_value, ensure_ascii=False)
                sqlitefun.sqlite_update_template(panel_model,panel_json)
                self.update_model = panel_model
                self.update_model_json = panel_json
                self.type=1
                self.send_main_Signal()
            else:
                QMessageBox.question(self, "错误提示", "输入的组件排布格式错误，请重新输入组件排布，格式为 2x10", QMessageBox.Yes)
    def btn_delede_model_Clicked(self):
        panel_model = self.panel_model.currentText()  # 组件型号
        use_model_name = sqlitefun.sqlite_select_use_template(panel_model)
        delete_use_model = False
        if len(use_model_name)>0:
            if use_model_name[0][0] == "":
                delete_use_model = True
            else:
                QMessageBox.question(self, "提示", panel_model+"模板在此项目或者其他项目中已被使用，不能删除此模板", QMessageBox.Yes)
        else:
            delete_use_model = True

        if delete_use_model:
            reply = QMessageBox.question(self, "提示", "你确定要删除模板:" + panel_model + "?",
                                         QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.Yes:
                sqlitefun.sqlite_delete_template(panel_model)
                reply_ok = QMessageBox.question(self, "提示", "删除模板:" + panel_model + "成功！",
                                                QMessageBox.Yes)
                if reply_ok == QMessageBox.Yes:
                    self.update_model = panel_model
                    self.update_model_json = ""
                    self.type = 2
                    self.send_main_Signal()
                    self.close()

    def send_main_Signal(self):
        lists = ["Update_Panel_Model_Signal",self.update_model,self.update_model_json,self.type]
        self.Update_Panel_Model_Signal.emit(lists)  # 发射信号


if __name__ =='__main__':
    app = QApplication(sys.argv)

    myWin = MyRGISMain()
    myWin.show()   

    sys.exit(app.exec_()) 