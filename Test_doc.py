from docx import Document
from docx.shared import Inches,RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

document.add_picture('Test_gps.jpg', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)
rgbColor = RGBColor(0,255,0)
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc
    tmp = "0,255,0"
    rgb = tmp.split(',')  # 将RGB格式划分开来
    colorStr = ''
    for i in rgb:
        num = int(i)  # 将str转int
        # 将R、G、B分别转化为16进制拼接转换并大写
        colorStr += str(hex(num))[-2:].replace('x', '0').upper()
    print("颜色值",colorStr)
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value=colorStr))
    row_cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
    #setCellBackgroundColor(row_cells[2],)
document.add_page_break()

document.save('demo.docx')



